#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
medallia_report_final.html → medallia_report_final.docx
디자인 스펙 가이드 기반 변환 스크립트
"""

import sys
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re
import lxml.etree as etree

# ─────────────────────────────────────────────
# 폰트 설정 헬퍼
# ─────────────────────────────────────────────
BATANG = '바탕체'

def set_font(run, size_pt, bold=False, italic=False, underline=False, color=None):
    """run에 바탕체 + 크기 + 스타일 적용"""
    run.font.name = BATANG
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    # eastAsia/hAnsi/cs 모두 바탕체
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), BATANG)
    rFonts.set(qn('w:hAnsi'), BATANG)
    rFonts.set(qn('w:eastAsia'), BATANG)
    rFonts.set(qn('w:cs'), BATANG)
    return run


def set_para_spacing(para, before_pt=0, after_pt=0, line_spacing_rule=None):
    """단락 간격 설정"""
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def set_para_indent(para, left_cm=0):
    """단락 왼쪽 들여쓰기 설정"""
    pf = para.paragraph_format
    pf.left_indent = Cm(left_cm)


def add_run_with_style(para, text, size_pt=14, bold=False, italic=False,
                        underline=False, color=None):
    """단락에 run 추가 후 스타일 적용"""
    if not text:
        return None
    run = para.add_run(text)
    set_font(run, size_pt, bold=bold, italic=italic, underline=underline, color=color)
    return run


# ─────────────────────────────────────────────
# 인라인 태그 처리 (strong, b, br 등)
# ─────────────────────────────────────────────
def process_inline(para, node, size_pt=14, bold_override=False,
                   italic_override=False, color=None):
    """node의 직접 자식을 순서대로 run으로 변환"""
    if node is None:
        return

    for child in node.children:
        if isinstance(child, NavigableString):
            txt = str(child)
            if txt:
                add_run_with_style(para, txt, size_pt,
                                   bold=bold_override, italic=italic_override,
                                   color=color)
        elif isinstance(child, Tag):
            tag = child.name
            if tag == 'br':
                run = para.add_run('\n')
                set_font(run, size_pt)
            elif tag in ('strong', 'b'):
                process_inline(para, child, size_pt,
                               bold_override=True,
                               italic_override=italic_override,
                               color=color)
            elif tag == 'em':
                process_inline(para, child, size_pt,
                               bold_override=bold_override,
                               italic_override=True,
                               color=color)
            elif tag == 'span':
                # span style 에서 color, font-size 처리
                style_str = child.get('style', '')
                sp_color = color
                sp_size = size_pt
                if 'color:#555' in style_str or 'color: #555' in style_str:
                    sp_color = (0x55, 0x55, 0x55)
                elif 'color:#c00' in style_str or 'color: #c00' in style_str:
                    sp_color = (0xcc, 0x00, 0x00)
                elif 'color:#444' in style_str or 'color: #444' in style_str:
                    sp_color = (0x44, 0x44, 0x44)
                if 'font-size:12' in style_str or 'font-size: 12' in style_str:
                    sp_size = 12
                elif 'font-size:13' in style_str or 'font-size: 13' in style_str:
                    sp_size = 13
                process_inline(para, child, sp_size,
                               bold_override=bold_override,
                               italic_override=italic_override,
                               color=sp_color)
            elif tag == 'a':
                # 링크 — 텍스트만 추출
                process_inline(para, child, size_pt,
                               bold_override=bold_override,
                               italic_override=italic_override,
                               color=(0x1a, 0x5e, 0xa8))
            else:
                process_inline(para, child, size_pt,
                               bold_override=bold_override,
                               italic_override=italic_override,
                               color=color)


# ─────────────────────────────────────────────
# li 직접 텍스트 추출 (하위 ul 제외)
# ─────────────────────────────────────────────
def get_li_direct_children(li_tag):
    """li 태그에서 하위 ul/ol을 제외한 직접 자식 노드들만 반환"""
    result = []
    for child in li_tag.children:
        if isinstance(child, Tag) and child.name in ('ul', 'ol'):
            continue
        result.append(child)
    return result


def li_has_only_text(li_tag):
    """li 직접 자식에 텍스트/인라인 태그만 있는지 확인"""
    for child in li_tag.children:
        if isinstance(child, Tag) and child.name in ('ul', 'ol'):
            continue
        if isinstance(child, Tag) and child.name in ('div', 'p', 'table'):
            return False
    return True


def add_para_from_li(doc, li_tag, bullet_char, size_pt=14,
                     indent_cm=0, after_pt=9, bold=False):
    """li에서 직접 텍스트만 추출해 단락 생성"""
    para = doc.add_paragraph()
    set_para_spacing(para, after_pt=after_pt)
    set_para_indent(para, left_cm=indent_cm)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # bullet 문자 run
    bullet_run = para.add_run(bullet_char + ' ')
    set_font(bullet_run, size_pt, bold=bold)

    # 직접 자식 처리
    direct_children = get_li_direct_children(li_tag)
    for child in direct_children:
        if isinstance(child, NavigableString):
            txt = str(child)
            if txt:
                add_run_with_style(para, txt, size_pt, bold=bold)
        elif isinstance(child, Tag):
            tag = child.name
            if tag == 'br':
                run = para.add_run('\n')
                set_font(run, size_pt)
            elif tag in ('strong', 'b'):
                process_inline(para, child, size_pt, bold_override=True)
            elif tag == 'em':
                process_inline(para, child, size_pt, italic_override=True)
            else:
                process_inline(para, child, size_pt)
    return para


# ─────────────────────────────────────────────
# 표 스타일 헬퍼
# ─────────────────────────────────────────────
def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = [('top', top), ('bottom', bottom), ('left', left), ('right', right)]
    for side, val in sides:
        el = OxmlElement(f'w:{side}')
        if val is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def add_cell_para(cell, text, size_pt=13, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  color=None, italic=False, after_pt=2):
    """셀에 단락 추가"""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        # 기존 내용 지우기
        for run in para.runs:
            run.text = ''
    else:
        para = cell.add_paragraph()
    para.alignment = align
    set_para_spacing(para, after_pt=after_pt)
    return para


def process_cell_html(cell, td_tag, size_pt=13):
    """td 태그의 HTML을 셀 단락으로 변환"""
    para = cell.paragraphs[0]
    set_para_spacing(para, after_pt=2)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # td 클래스 확인
    classes = td_tag.get('class', [])
    is_center = 'c' in classes
    is_bold = 'b' in classes
    is_red = 'r' in classes
    if is_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    color = None
    if is_red:
        color = (0xb0, 0x00, 0x00)

    # td 내용 처리
    has_content = False
    for child in td_tag.children:
        if isinstance(child, NavigableString):
            txt = str(child).strip()
            if txt:
                add_run_with_style(para, txt, size_pt, bold=is_bold, color=color)
                has_content = True
        elif isinstance(child, Tag):
            tag = child.name
            if tag == 'br':
                run = para.add_run('\n')
                set_font(run, size_pt)
                has_content = True
            elif tag in ('strong', 'b'):
                process_inline(para, child, size_pt, bold_override=True, color=color)
                has_content = True
            elif tag == 'span':
                style_str = child.get('style', '')
                sp_color = color
                sp_size = size_pt
                if 'color:#555' in style_str:
                    sp_color = (0x55, 0x55, 0x55)
                elif 'color:#c00' in style_str:
                    sp_color = (0xcc, 0x00, 0x00)
                if 'font-size:12' in style_str:
                    sp_size = 12
                elif 'font-size:13' in style_str:
                    sp_size = 13
                process_inline(para, child, sp_size, bold_override=is_bold, color=sp_color)
                has_content = True
            else:
                process_inline(para, child, size_pt, bold_override=is_bold, color=color)
                has_content = True

    if not has_content:
        add_run_with_style(para, '', size_pt)


# ─────────────────────────────────────────────
# 일반 표(table.t) 처리
# ─────────────────────────────────────────────
def add_regular_table(doc, table_tag):
    """table.t → Word 표"""
    rows_data = []
    header_rows = []
    body_rows = []

    thead = table_tag.find('thead')
    tbody = table_tag.find('tbody')

    if thead:
        for tr in thead.find_all('tr', recursive=False):
            header_rows.append(tr)
    if tbody:
        for tr in tbody.find_all('tr', recursive=False):
            body_rows.append(tr)
    else:
        for tr in table_tag.find_all('tr', recursive=False):
            body_rows.append(tr)

    all_rows = header_rows + body_rows
    if not all_rows:
        return

    # 열 수 결정
    max_cols = 0
    for tr in all_rows:
        cols = len(tr.find_all(['td', 'th'], recursive=False))
        max_cols = max(max_cols, cols)

    if max_cols == 0:
        return

    tbl = doc.add_table(rows=len(all_rows), cols=max_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 표 너비 100% 설정
    tbl_pr = tbl._tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), '5000')
    tbl_w.set(qn('w:type'), 'pct')
    tbl_pr.append(tbl_w)

    for row_idx, tr in enumerate(all_rows):
        is_header = tr in header_rows
        is_hl_row = 'hl-row' in tr.get('class', [])
        cells_html = tr.find_all(['td', 'th'], recursive=False)

        for col_idx, cell_html in enumerate(cells_html):
            if col_idx >= max_cols:
                break
            cell = tbl.cell(row_idx, col_idx)

            if is_header:
                # 헤더: 검정 배경, 흰 글자
                set_cell_bg(cell, '1E1E1E')
                para = cell.paragraphs[0]
                set_para_spacing(para, after_pt=2)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                txt = cell_html.get_text(separator=' ', strip=True)
                add_run_with_style(para, txt, 13, bold=True,
                                   color=(0xFF, 0xFF, 0xFF))
            elif is_hl_row:
                set_cell_bg(cell, 'FEF9E5')
                process_cell_html(cell, cell_html, 13)
            else:
                process_cell_html(cell, cell_html, 13)

    return tbl


# ─────────────────────────────────────────────
# 요약박스(summary-tbl) 처리
# ─────────────────────────────────────────────
def add_summary_table(doc, stbl_tag):
    """summary-tbl → 1행3열 Word 표 (중간 셀만 테두리)"""
    s_main_td = stbl_tag.find('td', class_='s-main')
    if not s_main_td:
        return

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 표 너비 100%
    tbl_pr = tbl._tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), '5000')
    tbl_w.set(qn('w:type'), 'pct')
    tbl_pr.append(tbl_w)

    # 표 전체 테두리 없애기
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

    # 패딩 셀 (좌)
    left_cell = tbl.cell(0, 0)
    tc_pr_l = left_cell._tc.get_or_add_tcPr()
    tc_w_l = OxmlElement('w:tcW')
    tc_w_l.set(qn('w:w'), '150')
    tc_w_l.set(qn('w:type'), 'pct')
    tc_pr_l.append(tc_w_l)

    # 패딩 셀 (우)
    right_cell = tbl.cell(0, 2)
    tc_pr_r = right_cell._tc.get_or_add_tcPr()
    tc_w_r = OxmlElement('w:tcW')
    tc_w_r.set(qn('w:w'), '150')
    tc_w_r.set(qn('w:type'), 'pct')
    tc_pr_r.append(tc_w_r)

    # 중간 셀 — 테두리 있음
    mid_cell = tbl.cell(0, 1)
    tc = mid_cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        tc_borders.append(el)
    tc_pr.append(tc_borders)

    # 중간 셀 내용
    para = mid_cell.paragraphs[0]
    set_para_spacing(para, after_pt=6)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    process_inline(para, s_main_td, size_pt=14)

    return tbl


# ─────────────────────────────────────────────
# 페이지 번호 바닥글 추가
# ─────────────────────────────────────────────
def add_page_number_footer(doc):
    """바닥글에 — {PAGE} — 형식 페이지 번호 추가"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # 기존 단락 활용
    if footer.paragraphs:
        para = footer.paragraphs[0]
        # 기존 내용 지우기
        for child in list(para._p):
            para._p.remove(child)
    else:
        para = footer.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, after_pt=0)

    # "— " 텍스트
    run1 = para.add_run('— ')
    set_font(run1, 12)

    # 페이지 번호 필드
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run_fld = para.add_run()
    run_fld._r.append(fld)
    set_font(run_fld, 12)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run_instr = para.add_run()
    run_instr._r.append(instrText)
    set_font(run_instr, 12)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_end = para.add_run()
    run_end._r.append(fld_end)
    set_font(run_end, 12)

    # " —" 텍스트
    run2 = para.add_run(' —')
    set_font(run2, 12)


# ─────────────────────────────────────────────
# 페이지 나누기 삽입
# ─────────────────────────────────────────────
def add_page_break(doc):
    """페이지 나누기 단락 추가"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    set_para_spacing(para, after_pt=0)


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


# ─────────────────────────────────────────────
# ol-row (목차 항목) 처리
# ─────────────────────────────────────────────
def process_ol_row(doc, row_div):
    """목차 ol-row 처리"""
    ol_num = row_div.find(class_='ol-num')
    ol_title = row_div.find(class_='ol-title')
    ol_subs = row_div.find(class_='ol-subs')

    # 번호 + 제목을 한 줄에
    para = doc.add_paragraph()
    set_para_spacing(para, after_pt=4)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if ol_num:
        num_text = ol_num.get_text(strip=True)
        add_run_with_style(para, num_text + '  ', 14, bold=True)

    if ol_title:
        title_text = ol_title.get_text(strip=True)
        add_run_with_style(para, title_text, 14, bold=True)

    # 소항목들
    if ol_subs:
        for span in ol_subs.find_all('span'):
            sub_para = doc.add_paragraph()
            set_para_spacing(sub_para, after_pt=3)
            set_para_indent(sub_para, left_cm=1.5)
            sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            txt = span.get_text(strip=True)
            add_run_with_style(sub_para, '□ ' + txt, 13)


# ─────────────────────────────────────────────
# ul 처리 (dash / dot / circ)
# ─────────────────────────────────────────────
def process_ul(doc, ul_tag, level=0):
    """ul.dash / ul.dot / ul.circ 처리 (중첩 지원)"""
    classes = ul_tag.get('class', [])

    # inline style에서 padding-left 확인 (override)
    style_str = ul_tag.get('style', '')
    has_extra_indent = 'padding-left:5em' in style_str

    if 'dash' in classes:
        bullet = '-'
        indent_cm = 2.0 if not has_extra_indent else 3.5
        size_pt = 14
        after_pt = 9
    elif 'dot' in classes:
        bullet = '·'
        indent_cm = 3.5
        size_pt = 14
        after_pt = 9
    elif 'circ' in classes:
        bullet = '○'
        indent_cm = 2.0
        size_pt = 14
        after_pt = 9
    else:
        bullet = '-'
        indent_cm = 1.5
        size_pt = 14
        after_pt = 9

    for li in ul_tag.find_all('li', recursive=False):
        # li 직접 텍스트 단락
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=after_pt)
        set_para_indent(para, left_cm=indent_cm)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        bullet_run = para.add_run(bullet + ' ')
        set_font(bullet_run, size_pt)

        # 직접 자식 처리 (하위 ul 제외)
        direct_children = get_li_direct_children(li)
        for child in direct_children:
            if isinstance(child, NavigableString):
                txt = str(child)
                if txt:
                    add_run_with_style(para, txt, size_pt)
            elif isinstance(child, Tag):
                tag = child.name
                if tag == 'br':
                    run = para.add_run('\n')
                    set_font(run, size_pt)
                elif tag in ('strong', 'b'):
                    process_inline(para, child, size_pt, bold_override=True)
                elif tag == 'em':
                    process_inline(para, child, size_pt, italic_override=True)
                else:
                    process_inline(para, child, size_pt)

        # 하위 ul 처리
        for sub_ul in li.find_all('ul', recursive=False):
            process_ul(doc, sub_ul, level=level + 1)


# ─────────────────────────────────────────────
# 문서 초기 설정
# ─────────────────────────────────────────────
def setup_document():
    """문서 기본 설정"""
    doc = Document()

    # A4 용지 설정
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.footer_distance = Mm(8)

    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = BATANG
    font.size = Pt(14)
    # eastAsia 설정
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        style.element.append(rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), BATANG)
    rFonts.set(qn('w:hAnsi'), BATANG)
    rFonts.set(qn('w:eastAsia'), BATANG)
    rFonts.set(qn('w:cs'), BATANG)

    return doc


# ─────────────────────────────────────────────
# 페이지 내 요소 처리
# ─────────────────────────────────────────────
def process_page_element(doc, element, is_first_page=False):
    """doc-page 내 개별 요소 처리"""
    if not isinstance(element, Tag):
        return

    tag = element.name
    classes = element.get('class', []) if element.get('class') else []
    style_str = element.get('style', '') or ''

    # .page-no → 무시 (바닥글로 처리)
    if 'page-no' in classes:
        return

    # .doc-main-title
    if 'doc-main-title' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        txt = element.get_text(strip=True)
        add_run_with_style(para, txt, 20, bold=True, underline=True)
        return

    # .doc-hd-title
    if 'doc-hd-title' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=4)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        txt = element.get_text(strip=True)
        add_run_with_style(para, txt, 14, bold=True)
        return

    # .doc-hd-meta
    if 'doc-hd-meta' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        txt = element.get_text(strip=True)
        add_run_with_style(para, txt, 14)
        return

    # .doc-hd-bar → 구분선
    if 'doc-hd-bar' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=2)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top_bdr = OxmlElement('w:top')
        top_bdr.set(qn('w:val'), 'single')
        top_bdr.set(qn('w:sz'), '20')
        top_bdr.set(qn('w:color'), '000000')
        pBdr.append(top_bdr)
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single')
        bottom_bdr.set(qn('w:sz'), '6')
        bottom_bdr.set(qn('w:color'), '000000')
        pBdr.append(bottom_bdr)
        pPr.append(pBdr)
        return

    # 날짜/소속 — text-align:right div
    if tag == 'div' and 'text-align:right' in style_str:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        process_inline(para, element, 14)
        return

    # 챕터 제목 (별첨 특수 제목 포함)
    if tag == 'div' and 'text-align:center' in style_str and ('15px' in style_str or '15' in style_str):
        # 별첨 페이지 제목 【 별첨 N 】...
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        txt = element.get_text(strip=True)
        add_run_with_style(para, txt, 15, bold=True)
        return

    # .ch → Lv1 챕터
    if 'ch' in classes or 'app-ch' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=12)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        process_inline(para, element, 14, bold_override=True)
        return

    # .sq → Lv2 □ 소제목
    if 'sq' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=16)
        set_para_indent(para, left_cm=0.5)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run_with_style(para, '□ ', 14)
        process_inline(para, element, 14)
        return

    # .app-sub → 밑줄 소제목
    if 'app-sub' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=12)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        txt = element.get_text(strip=True)
        add_run_with_style(para, txt, 14, bold=True, underline=True)
        return

    # .fn → Lv5 ※ 주석 (10pt)
    if 'fn' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        set_para_indent(para, left_cm=3.0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run_with_style(para, '※ ', 10, bold=True)
        process_inline(para, element, 10)
        return

    # .body-p
    if 'body-p' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        set_para_indent(para, left_cm=2.0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        process_inline(para, element, 14)
        return

    # .quote
    if 'quote' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        set_para_indent(para, left_cm=1.0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 인용문 텍스트 (quote-src 제외)
        quote_src = element.find(class_='quote-src')
        for child in element.children:
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if txt:
                    add_run_with_style(para, txt, 13, italic=True,
                                       color=(0x33, 0x33, 0x33))
            elif isinstance(child, Tag):
                if 'quote-src' in (child.get('class') or []):
                    continue
                process_inline(para, child, 13, italic_override=True,
                               color=(0x33, 0x33, 0x33))

        if quote_src:
            src_para = doc.add_paragraph()
            set_para_spacing(src_para, after_pt=9)
            set_para_indent(src_para, left_cm=1.0)
            src_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            txt = quote_src.get_text(strip=True)
            add_run_with_style(src_para, txt, 12, italic=True,
                               color=(0x66, 0x66, 0x66))
        return

    # .box / .box-key
    if 'box' in classes or 'box-key' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 회색 배경 박스 — 테두리 단락
        pPr = para._p.get_or_add_pPr()
        # 배경 음영 설정
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        # 테두리 설정
        pBdr = OxmlElement('w:pBdr')
        for side in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), '888888')
            pBdr.append(el)
        pPr.append(pBdr)
        process_inline(para, element, 13)
        return

    # .arr → 화살표 소결론
    if 'arr' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        a_span = element.find(class_='a')
        if a_span:
            add_run_with_style(para, '→ ', 14, bold=True)
        # 나머지 텍스트
        for child in element.children:
            if isinstance(child, Tag) and 'a' in (child.get('class') or []):
                continue
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if txt:
                    add_run_with_style(para, txt, 14)
            elif isinstance(child, Tag):
                process_inline(para, child, 14)
        return

    # summary-tbl
    if 'summary-tbl' in classes:
        add_summary_table(doc, element)
        return

    # table.t → 일반 표
    if tag == 'table' and 't' in classes:
        add_regular_table(doc, element)
        return

    # ul.dash / ul.dot / ul.circ
    if tag == 'ul':
        process_ul(doc, element)
        return

    # ol-row (목차 항목)
    if 'ol-row' in classes:
        process_ol_row(doc, element)
        return

    # hr.sep / hr.sep-dash → 구분선
    if tag == 'hr':
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=6)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top_bdr = OxmlElement('w:top')
        top_bdr.set(qn('w:val'), 'single')
        top_bdr.set(qn('w:sz'), '4')
        top_bdr.set(qn('w:color'), 'CCCCCC')
        pBdr.append(top_bdr)
        pPr.append(pBdr)
        return

    # SVG → 텍스트 대체
    if tag == 'svg':
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(para,
            '[주요 재무 지표 비교 차트 — 이자비용($300M)이 EBITDA($200M) 초과, 구조적 적자 확인]',
            12, italic=True, color=(0x44, 0x44, 0x44))
        return

    # p / div — 일반 텍스트 처리
    if tag in ('p', 'div'):
        # font-weight:700 div → Bold 소제목
        if 'font-weight:700' in style_str or 'font-weight: 700' in style_str:
            para = doc.add_paragraph()
            set_para_spacing(para, after_pt=9)

            # padding-left 확인
            indent_cm = 0.0
            if 'padding-left:3em' in style_str:
                indent_cm = 2.0
            set_para_indent(para, left_cm=indent_cm)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            process_inline(para, element, 14, bold_override=True)
            return

        # text-align:center
        if 'text-align:center' in style_str:
            para = doc.add_paragraph()
            set_para_spacing(para, after_pt=6)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # font-size 확인
            size = 12
            if 'font-size:12px' in style_str or 'font-size: 12px' in style_str:
                size = 12
            elif 'font-size:11.5px' in style_str or 'font-size: 11.5px' in style_str:
                size = 10
            elif 'font-size:14px' in style_str:
                size = 14

            color = None
            if 'color:#444' in style_str:
                color = (0x44, 0x44, 0x44)
            elif 'color:#c00' in style_str:
                color = (0xcc, 0x00, 0x00)
            elif 'color:#777' in style_str:
                color = (0x77, 0x77, 0x77)

            bold = 'font-weight:700' in style_str or 'font-weight: 700' in style_str
            process_inline(para, element, size, bold_override=bold, color=color)
            return

        # 일반 div/p
        txt = element.get_text(strip=True)
        if not txt:
            return

        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        size = 14
        color = None
        if 'font-size:12.5px' in style_str or 'font-size: 12.5px' in style_str:
            size = 12
        elif 'font-size:14px' in style_str:
            size = 14

        if 'color:#333' in style_str:
            color = (0x33, 0x33, 0x33)

        bold = 'font-weight:700' in style_str or 'font-weight: 700' in style_str
        process_inline(para, element, size, bold_override=bold, color=color)
        return

    # ref-cat
    if 'ref-cat' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=6)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        txt = element.get_text(strip=True)
        add_run_with_style(para, txt, 12, bold=True,
                           color=(0x33, 0x33, 0x33))
        return

    # gap → 빈 단락
    if 'gap' in classes:
        para = doc.add_paragraph()
        set_para_spacing(para, after_pt=9)
        return


# ─────────────────────────────────────────────
# 메인 변환 로직
# ─────────────────────────────────────────────
def convert_html_to_docx(html_path, output_path):
    print(f"HTML 파일 로드: {html_path}")

    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    pages = soup.find_all('div', class_='doc-page')
    print(f"총 {len(pages)}개 페이지 감지")

    doc = setup_document()

    # 첫 페이지는 목차 (페이지 번호 없음)
    # 본문 페이지부터 바닥글 페이지 번호 적용
    add_page_number_footer(doc)

    for page_idx, page in enumerate(pages):
        if page_idx > 0:
            # 페이지 나누기
            para = doc.add_paragraph()
            run = para.add_run()
            from docx.oxml import OxmlElement as _OE
            br = _OE('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)
            set_para_spacing(para, after_pt=0)

        # 페이지 내 모든 직접 자식 요소 처리
        for child in page.children:
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if txt:
                    para = doc.add_paragraph()
                    set_para_spacing(para, after_pt=6)
                    add_run_with_style(para, txt, 14)
            elif isinstance(child, Tag):
                tag = child.name
                classes = child.get('class') or []

                # 이름 없는 공백 태그 건너뛰기
                if not child.get_text(strip=True) and tag not in ('hr', 'svg', 'table'):
                    continue

                process_page_element(doc, child, is_first_page=(page_idx == 0))

    # 문서 끝 — 以 上 —
    para = doc.add_paragraph()
    set_para_spacing(para, after_pt=12)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run_with_style(para, '— 以 上 —', 14)

    doc.save(output_path)
    print(f"저장 완료: {output_path}")
    return output_path


# ─────────────────────────────────────────────
# 실행
# ─────────────────────────────────────────────
if __name__ == '__main__':
    import os
    HTML_PATH = '/home/ubuntu/medallia_report_final.html'
    OUTPUT_PATH = '/home/ubuntu/files/medallia_report_final.docx'

    try:
        result = convert_html_to_docx(HTML_PATH, OUTPUT_PATH)
        size = os.path.getsize(result)
        print(f"\n✓ 변환 성공!")
        print(f"  파일: {result}")
        print(f"  크기: {size:,} bytes ({size/1024:.1f} KB)")
    except Exception as e:
        import traceback
        print(f"\n✗ 오류 발생:")
        traceback.print_exc()
        sys.exit(1)
