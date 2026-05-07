#!/usr/bin/env python3
"""medallia_report_final.html → medallia_report_final.docx 변환 스크립트"""

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re

# ── 헬퍼 ──────────────────────────────────────────────────────────────────────

def set_font(run, name='맑은 고딕', size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_font(para, name='맑은 고딕', size=10, space_before=0, space_after=2,
              line_spacing=None, left_indent=None, align=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if line_spacing:
        from docx.shared import Pt as Pt2
        fmt.line_spacing = Pt2(line_spacing)
    if left_indent is not None:
        fmt.left_indent = Cm(left_indent)
    if align:
        para.alignment = align


def add_hr(doc, thick=False):
    """수평선 단락 삽입"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12' if thick else '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000' if thick else 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


def get_text_clean(el):
    """태그 텍스트를 깔끔하게 추출 (page-badge, page-no 제외)"""
    for badge in el.find_all(class_=['page-badge', 'page-no']):
        badge.decompose()
    txt = el.get_text(' ', strip=True)
    txt = re.sub(r'\s+', ' ', txt)
    return txt


def inline_run(para, text, bold=False, italic=False,
               size=10, color=None, name='맑은 고딕'):
    if not text:
        return
    run = para.add_run(text)
    set_font(run, name=name, size=size, bold=bold, italic=italic, color=color)


def render_inline(para, el, default_size=10):
    """el 내부의 인라인 요소를 para에 런으로 렌더"""
    if el is None:
        return
    for node in el.children:
        if isinstance(node, NavigableString):
            txt = str(node).replace('\xa0', ' ')
            if txt.strip() or txt:
                inline_run(para, txt, size=default_size)
        elif isinstance(node, Tag):
            tag = node.name
            cls = node.get('class', [])
            style = node.get('style', '')
            txt = node.get_text(' ', strip=True).replace('\xa0', ' ')
            if not txt:
                continue
            bold = tag in ('strong', 'b') or 'font-weight:700' in style or 'font-weight: 700' in style
            italic = tag in ('em', 'i') or 'font-style:italic' in style
            color = None
            if 'color:#b00' in style or 'color: #b00' in style:
                color = (0xbb, 0x00, 0x00)
            size = default_size
            if 'font-size:12px' in style or 'font-size: 12px' in style:
                size = 9
            if 'font-size:12.5px' in style or 'font-size:11.5px' in style:
                size = 9
            if tag == 'br':
                inline_run(para, '\n', size=size)
            elif tag == 'span' and 'page-badge' in cls:
                pass  # 뱃지 무시
            elif tag == 'a':
                inline_run(para, txt, size=9, color=(0x1a, 0x5e, 0xa8))
            else:
                inline_run(para, txt, bold=bold, italic=italic, size=size, color=color)


# ── 셀 스타일 ──────────────────────────────────────────────────────────────────

def style_cell(cell, bg=None, bold=False, align='left', color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            if bold:
                run.font.bold = True
            if color:
                run.font.color.rgb = RGBColor(*color)
        if align == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)


def add_table_from_html(doc, tbl_el):
    """table.t → docx 표"""
    rows = tbl_el.find_all('tr')
    if not rows:
        return

    # 열 수 계산
    max_cols = 0
    for row in rows:
        cols = sum(int(c.get('colspan', 1)) for c in row.find_all(['th', 'td']))
        max_cols = max(max_cols, cols)

    if max_cols == 0:
        return

    tbl = doc.add_table(rows=0, cols=max_cols)
    tbl.style = 'Table Grid'

    for row_el in rows:
        row = tbl.add_row()
        cells = row_el.find_all(['th', 'td'])
        col_idx = 0
        for cell_el in cells:
            if col_idx >= max_cols:
                break
            cell = row.cells[col_idx]
            tag = cell_el.name
            cls = cell_el.get('class', [])
            style = cell_el.get('style', '')
            is_header = (tag == 'th')
            is_bold = ('b' in cls) or is_header
            is_red = 'r' in cls
            is_center = ('c' in cls) or is_header
            is_hl = 'hl' in cls or 'hl-row' in row_el.get('class', [])

            # 셀 텍스트
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            render_inline(para, cell_el, default_size=9)

            for run in para.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(9)
                if is_header:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                elif is_bold:
                    run.font.bold = True
                if is_red:
                    run.font.color.rgb = RGBColor(0xbb, 0x00, 0x00)
            if is_center:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 배경색
            fill = None
            if is_header:
                fill = '1e1e1e'
            elif is_hl:
                fill = 'FEF9E5'
            if fill:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                tcPr.append(shd)

            # colspan 처리
            colspan = int(cell_el.get('colspan', 1))
            if colspan > 1:
                for merge_idx in range(1, colspan):
                    if col_idx + merge_idx < max_cols:
                        cell = cell.merge(row.cells[col_idx + merge_idx])

            col_idx += colspan


# ── 블록 렌더러 ───────────────────────────────────────────────────────────────

def process_li(doc, li_el, level=0, bullet_char='-'):
    """li 재귀 처리"""
    indent = Cm(0.6 * (level + 1))
    # 직계 텍스트/인라인만 추출 (자식 ul 제외)
    direct_text_el = copy.copy(li_el)
    for sub in direct_text_el.find_all(['ul', 'ol']):
        sub.decompose()

    raw = direct_text_el.get_text(' ', strip=True).replace('\xa0', ' ')
    raw = re.sub(r'\s+', ' ', raw).strip()
    if not raw:
        raw = ''

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Pt(-12)

    # 불릿 문자 + 내용
    prefix_run = p.add_run(bullet_char + ' ')
    set_font(prefix_run, size=10)
    # 인라인 렌더
    render_inline(p, direct_text_el, default_size=10)

    # 하위 ul 처리
    for sub_ul in li_el.find_all('ul', recursive=False):
        sub_cls = sub_ul.get('class', [])
        if 'dot' in sub_cls:
            sub_bullet = '·'
        elif 'circ' in sub_cls:
            sub_bullet = '○'
        else:
            sub_bullet = '-'
        for sub_li in sub_ul.find_all('li', recursive=False):
            process_li(doc, sub_li, level=level + 1, bullet_char=sub_bullet)


def process_block(doc, el):
    """하나의 블록 요소 처리"""
    if not isinstance(el, Tag):
        return

    tag = el.name
    cls = el.get('class', [])

    # 무시
    if 'page-no' in cls or 'page-badge' in cls:
        return

    # ── 문서 헤더 타이틀
    if 'doc-hd-title' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        txt = get_text_clean(el)
        r = p.add_run(txt)
        set_font(r, size=12, bold=True)
        return

    # ── 수평선 바
    if 'doc-hd-bar' in cls or tag in ('hr',):
        add_hr(doc, thick=('doc-hd-bar' in cls))
        return

    # ── 메타 (날짜/부서)
    if 'doc-hd-meta' in cls:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(el.get_text(' ', strip=True).replace('\xa0', ' '))
        set_font(r, size=9)
        return

    # ── 리드 문단
    if 'doc-lead' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(0.3)
        txt = el.get_text(' ', strip=True).replace('\xa0', ' ')
        txt = re.sub(r'\s+', ' ', txt).strip()
        r = p.add_run(txt)
        set_font(r, size=10, italic=False)
        # 아래 구분선
        fmt = p.paragraph_format
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'DDDDDD')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return

    # ── 챕터 제목
    if 'ch' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(el.get_text(' ', strip=True).replace('\xa0', ' '))
        set_font(r, size=11, bold=True)
        return

    # ── □ 소제목
    if 'sq' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(' □ ' + el.get_text(' ', strip=True).replace('\xa0', ' '))
        set_font(r, size=10, bold=True)
        return

    # ── 본문 들여쓰기 단락
    if 'body-p' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.7)
        render_inline(p, el, default_size=10)
        return

    # ── 각주 fn
    if 'fn' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.2)
        r_prefix = p.add_run('※ ')
        set_font(r_prefix, size=9, bold=True)
        txt = el.get_text(' ', strip=True).replace('\xa0', ' ')
        txt = re.sub(r'\s+', ' ', txt).strip()
        r = p.add_run(txt)
        set_font(r, size=9, color=(0x33, 0x33, 0x33))
        return

    # ── 인용 quote
    if 'quote' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        # 왼쪽 세로선
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '999999')
        pBdr.append(left)
        pPr.append(pBdr)
        # 본문
        src_el = el.find(class_='quote-src')
        src_text = ''
        if src_el:
            src_text = src_el.get_text(' ', strip=True)
            src_el.decompose()
        main_txt = el.get_text(' ', strip=True)
        main_txt = re.sub(r'\s+', ' ', main_txt).strip()
        r = p.add_run(main_txt)
        set_font(r, size=9, italic=True, color=(0x33, 0x33, 0x33))
        if src_text:
            r2 = p.add_run('\n' + src_text)
            set_font(r2, size=8.5, italic=False, color=(0x66, 0x66, 0x66))
        return

    # ── 박스 box / box-key
    if 'box' in cls or 'box-key' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        # 박스 테두리
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'), 'single')
            bd.set(qn('w:sz'), '6' if 'box-key' in cls and side == 'left' else '4')
            bd.set(qn('w:space'), '4')
            bd.set(qn('w:color'), '888888' if 'box-key' in cls and side == 'left' else 'AAAAAA')
            pBdr.append(bd)
        pPr.append(pBdr)
        render_inline(p, el, default_size=9.5)
        return

    # ── 화살표 소결 arr
    if 'arr' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(3)
        r_arrow = p.add_run('→ ')
        set_font(r_arrow, size=10, bold=True)
        # .a 스팬(→)을 제거하고 나머지 텍스트/인라인 렌더
        el_copy = copy.copy(el)
        for a_span in el_copy.find_all(class_='a'):
            a_span.decompose()
        render_inline(p, el_copy, default_size=10)
        return

    # ── 리스트
    if tag == 'ul':
        ul_cls = el.get('class', [])
        if 'dot' in ul_cls:
            bullet = '·'
        elif 'circ' in ul_cls:
            bullet = '○'
        else:
            bullet = '-'
        for li in el.find_all('li', recursive=False):
            process_li(doc, li, level=0, bullet_char=bullet)
        return

    # ── 표
    if tag == 'table':
        add_table_from_html(doc, el)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return

    # ── 구분선
    if tag == 'hr':
        add_hr(doc)
        return

    # ── 별첨 제목 app-ch / app-sub
    if 'app-ch' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        txt = el.get_text(' ', strip=True)
        r = p.add_run(txt)
        set_font(r, size=11, bold=True)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '333333')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return

    if 'app-sub' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        txt = el.get_text(' ', strip=True)
        r = p.add_run(txt)
        set_font(r, size=10, bold=True)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return

    # ── ol-row (목차 행)
    if 'ol-row' in cls:
        num_el = el.find(class_='ol-num')
        title_el = el.find(class_='ol-title')
        subs_el = el.find(class_='ol-subs')
        num = num_el.get_text(strip=True) if num_el else ''
        title = title_el.get_text(strip=True) if title_el else ''
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(num + '  ')
        set_font(r1, size=10, bold=True)
        r2 = p.add_run(title)
        set_font(r2, size=10, bold=True)
        if subs_el:
            for span in subs_el.find_all('span'):
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(1)
                sp.paragraph_format.space_after = Pt(1)
                sp.paragraph_format.left_indent = Cm(1.5)
                r = sp.add_run('□ ' + span.get_text(strip=True))
                set_font(r, size=9, color=(0x44, 0x44, 0x44))
        return

    # ── ref-cat
    if 'ref-cat' in cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(el.get_text(strip=True))
        set_font(r, size=9.5, bold=True, color=(0x33, 0x33, 0x33))
        return

    # ── 일반 p 태그 (인라인 콘텐츠)
    if tag == 'p':
        style = el.get('style', '')
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if 'text-align:center' in style or 'text-align: center' in style:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        render_inline(p, el, default_size=10)
        return

    # ── div 중 style="text-align:center" (별첨 제목)
    if tag == 'div':
        style = el.get('style', '')
        if 'text-align:center' in style or 'text-align: center' in style:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            txt = el.get_text(' ', strip=True)
            r = p.add_run(txt)
            set_font(r, size=11, bold=True)
            return
        # 하위 블록 재귀
        for child in el.children:
            if isinstance(child, Tag):
                process_block(doc, child)
        return

    # ── SVG 차트는 건너뜀 (텍스트만 표현 불가)
    if tag == 'svg':
        p = doc.add_paragraph()
        r = p.add_run('[※ 차트 이미지는 Word 변환에서 생략됩니다]')
        set_font(r, size=9, color=(0x88, 0x88, 0x88), italic=True)
        return

    # ── 나머지: 텍스트가 있으면 단순 단락
    txt = el.get_text(' ', strip=True)
    if txt:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        render_inline(p, el, default_size=10)


# ── 메인 ──────────────────────────────────────────────────────────────────────

def convert():
    import docx as _docx_mod  # noqa
    global docx
    docx = _docx_mod

    with open('/home/ubuntu/medallia_report_final.html', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    doc = Document()

    # 페이지 여백 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)

    # 기본 스타일 초기화
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.font.size = Pt(10)

    pages = soup.find_all(class_='doc-page')
    print(f'총 {len(pages)}페이지 변환 시작...')

    for page_num, page in enumerate(pages, 1):
        if page_num > 1:
            # 페이지 구분선
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)

        print(f'  페이지 {page_num} 처리 중...')
        for child in page.children:
            if isinstance(child, Tag):
                process_block(doc, child)

    out_path = '/home/ubuntu/medallia_report_final.docx'
    doc.save(out_path)
    print(f'\n완료: {out_path}')


if __name__ == '__main__':
    convert()
