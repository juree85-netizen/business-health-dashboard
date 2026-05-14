#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Medallia 사례보고서 — analysis-report-standard.md 기준 Word 문서 생성
"""
import sys, os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BATANG = '바탕체'
PT_BODY  = 14
PT_TITLE = 20
PT_PGNUM = 12

# ── 폰트 헬퍼 ──────────────────────────────────────────────────────────────────

def set_font(run, size=PT_BODY, bold=False, italic=False, underline=False, color=None):
    run.font.name = BATANG
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rf.set(qn(a), BATANG)

def set_pf(para, before=0, after=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.alignment = align

def add_run(para, text, **kw):
    r = para.add_run(text)
    set_font(r, **kw)
    return r

# ── 문서 초기 설정 ────────────────────────────────────────────────────────────

def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width      = Mm(210)
    sec.page_height     = Mm(297)
    sec.top_margin      = Mm(20)
    sec.bottom_margin   = Mm(20)
    sec.left_margin     = Mm(20)
    sec.right_margin    = Mm(20)
    sec.footer_distance = Mm(8)
    # Normal 스타일 — 바탕체 14pt
    sty = doc.styles['Normal']
    sty.font.name = BATANG
    sty.font.size = Pt(PT_BODY)
    el = sty.element
    rpr = el.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr'); el.append(rpr)
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rf.set(qn(a), BATANG)
    return doc

# ── 바닥글 페이지 번호 ─────────────────────────────────────────────────────────

def add_footer(doc):
    sec = doc.sections[0]
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    for ch in list(para._p): para._p.remove(ch)
    set_pf(para, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(para, '- ', size=PT_PGNUM)
    # PAGE 필드
    for ftype in ('begin', None, 'end'):
        r = para.add_run()
        set_font(r, size=PT_PGNUM)
        if ftype in ('begin','end'):
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), ftype)
            r._r.append(fc)
        else:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = ' PAGE '
            r._r.append(it)
    add_run(para, ' -', size=PT_PGNUM)

# ── 요약 표 (1행 3열) ──────────────────────────────────────────────────────────

def add_summary_table(doc, lines):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tpr = tbl._tbl.find(qn('w:tblPr'))
    if tpr is None:
        tpr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tpr)
    # 전체 테두리 제거
    tb = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'),'nil'); tb.append(e)
    tpr.append(tb)
    # 자동 맞춤 해제
    tl = OxmlElement('w:tblLayout'); tl.set(qn('w:type'),'fixed'); tpr.append(tl)

    def cell_w(cell, dxa):
        cp = cell._tc.get_or_add_tcPr()
        cw = OxmlElement('w:tcW')
        cw.set(qn('w:w'), str(dxa)); cw.set(qn('w:type'),'dxa')
        cp.append(cw)

    def cell_margin(cell, l=0, r=0, t=0, b=0):
        cp = cell._tc.get_or_add_tcPr()
        cm = OxmlElement('w:tcMar')
        for s,v in (('left',l),('right',r),('top',t),('bottom',b)):
            e = OxmlElement(f'w:{s}')
            e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa')
            cm.append(e)
        cp.append(cm)

    def cell_borders(cell, has, nil):
        cp = cell._tc.get_or_add_tcPr()
        cb = OxmlElement('w:tcBorders')
        for s in has:
            e = OxmlElement(f'w:{s}')
            e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
            e.set(qn('w:color'),'000000'); cb.append(e)
        for s in nil:
            e = OxmlElement(f'w:{s}'); e.set(qn('w:val'),'nil'); cb.append(e)
        cp.append(cb)

    NARROW = 140    # 7pt × 20 = 140 twips
    CWIDE  = 9353   # 165mm × 56.69 ≈ 9353 twips
    LC = tbl.cell(0,0); MC = tbl.cell(0,1); RC = tbl.cell(0,2)
    cell_w(LC, NARROW); cell_w(MC, CWIDE); cell_w(RC, NARROW)
    cell_margin(LC); cell_margin(RC)
    cell_margin(MC, l=28, r=28)   # 좌우 1.4pt
    cell_borders(LC, ['top','left','bottom'], ['right'])
    cell_borders(MC, [], ['top','left','bottom','right'])
    cell_borders(RC, ['top','right','bottom'], ['left'])
    # 좌우 셀 빈 단락
    for cell in (LC, RC):
        set_pf(cell.paragraphs[0])
    # 가운데 셀 — 3줄
    for i, line in enumerate(lines[:3]):
        p = MC.paragraphs[0] if i == 0 else MC.add_paragraph()
        set_pf(p, before=0, after=9)
        add_run(p, ' ' + line)   # 선행 Space 1칸

# ── 일반 표 ───────────────────────────────────────────────────────────────────

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tpr = tbl._tbl.find(qn('w:tblPr'))
    if tpr is None:
        tpr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tpr)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'),'5000'); tw.set(qn('w:type'),'pct'); tpr.append(tw)
    # 헤더
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cp = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'),'1E1E1E'); cp.append(shd)
        p = cell.paragraphs[0]
        set_pf(p, before=1, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, h, size=13, bold=True, color=(0xFF,0xFF,0xFF))
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[1+ri].cells[ci]
            p = cell.paragraphs[0]
            set_pf(p, before=1, after=1)
            add_run(p, str(val), size=13)

# ── 페이지 나누기 ─────────────────────────────────────────────────────────────

def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'),'page'); r._r.append(br)
    set_pf(p)

# ── 콘텐츠 정의 ───────────────────────────────────────────────────────────────

REPORT = [
    # ── 1. 보고 개요 ──
    ('lv1','1. 보고 개요'),
    ('lv2','분석 배경'),
    ('lv3',"Thoma Bravo는 Medallia를 $6.4B(멀티플 약 9배)에 인수하며 상장폐지 ('21.7月)"),
    ('lv3',"이후 美 연준 금리 인상에 따라 SOFR가 약 5.4%까지 상승하며 차입 이자 부담 급증 ('22~'25年)"),
    ('lv3',"EBITDA를 초과하는 이자 비용 발생으로 구조적 부실 상태에 직면, 채권단으로 경영권 이전 및 투자금 손실 발생 ('26.4月)"),
    ('lv2','분석 범위'),
    ('lv3',"분석 대상: Medallia / Thoma Bravo"),
    ('lv3',"분석 기간: '21年 Thoma Bravo의 Medallia 인수 ~ '26年 채권단으로 경영권 이전"),
    ('lv3',"재무 수치 출처: Bloomberg, OnlyCFO 분석 기반 추정치"),
    ('lv2','분석 목적'),
    ('lv3',"PE 기반 SaaS 차입매수(LBO) 방식의 실패 원인 분석"),
    ('lv3',"SaaS 기업가치 평가 기준 재점검 — ARR성장률 중심 평가 방식의 한계점 확인"),
    # ── 2. Medallia 개요 ──
    ('lv1','2. Medallia 개요'),
    ('lv2','기업 개요'),
    ('lv3',"'01年 설립, 미국 샌프란시스코 소재 CX 관리 플랫폼 전문 기업"),
    ('lv3',"CRM 솔루션 영역 중 CX Management(고객경험 분석 관리) 부문 內 2위"),
    ('lv4',"CRM은 Pre-Sales, Sales, Post-sales(Service/Support)로 구분되며, CX Management는 Post-sales 시장과 연관"),
    ('lv3',"오퍼링: 고객경험 피드백 수집/분석, NPS 측정, 직원경험 플랫폼, 옴니채널 고객 인텔리전스 제공"),
    ('lv3',"고객군: 금융, 항공, 호텔, 통신 등 글로벌 대기업"),
    ('lv3',"'15~'18年: VC, PE로부터 복수의 투자 유치, 기업가치 급성장"),
    ('lv3',"'19年: NYSE 상장, IPO 공모가 $21 → 상장 이후 $30~$40"),
    ('lv3',"'20年: 매출 $477M, 순손실 –$149M"),
    ('lv3',"'21年: Thoma Bravo에 인수되며 상장폐지"),
    # ── 3. Medallia 투자/인수 현황 ──
    ('lv1','3. Medallia 투자/인수 현황'),
    ('lv2',"Thoma Bravo의 Medallia 인수 ('21年)"),
    ('lv3',"당시 고성장 SaaS 프리미엄을 고려, EV/Revenue 멀티플을 약 9배로 적용"),
    ('lv5',"멀티플: 기업가치(EV)를 특정 재무지표로 나눈 값으로, 시장이 해당 기업 실적에 부여하는 가치 수준"),
    ('lv3',"자금 조달 구조"),
    ('table',
     ['구분','금액','조달처'],
     [
         ['투자금 (Equity)','$4.6B','Thoma Bravo 등'],
         ['변동금리 차입금 (Debt)','$1.8B','민간대출업자 (Blackstone 등)'],
         ['총 인수가격','$6.4B',''],
     ]),
    ('lv5',"LIBOR: 런던 은행 간 단기 대출 금리 (글로벌 금융 계약 기준금리, '23年 SOFR로 대체). 당시 0.75%"),
    ('lv3',"인수 당시 계획: 저금리 환경에서 낮은 비용으로 차입 매수, 영업현금흐름으로 이자 감당, 3~5년 내 매각 또는 IPO 재상장으로 투자금 회수"),
    ('lv3',"인수 판단 기준"),
    ('lv4',"Medallia는 ARR 성장률 23~26%의 고성장 기업으로 SaaS 차입매수(LBO)가 가능한 기업으로 판단"),
    ('lv4',"단, Rule of 40·잉여현금흐름 마진은 동종업계 하위 10% 수준임을 인지"),
    ('lv4',"당시 초저금리 환경으로 낮은 이자 부담, PIK 활용 가능하다고 판단"),
    ('lv2','인수 이후 상황'),
    ('lv3',"1) 美 연준 금리 인상으로 이자부담 급증 ('22年 이후)"),
    ('table',
     ['구분',"'21年","'23年","'26年(예상)"],
     [
         ['기준금리(SOFR)','0.75%','5.4%','3.7%'],
         ['연이자','$135M','$300M','$270M'],
     ]),
    ('lv4',"부채 규모 확대: PIK 적용으로 원금 $1.8B → 약 $2.2B (현금 미지급 이자가 원금에 누적 반영)"),
    ('lv5',"PIK 비중 점진 축소 ('21~'22年 100% → '22~'23年 50% → '23~'25年 35% → '25~'26年 0%)"),
    ('lv5',"Medallia를 통한 추가 M&A 관련 차입금($0.6B) 반영 시 총부채 약 $2.8B 수준"),
    ('lv3',"2) 매각, IPO 재진입 실패로 투자금 회수 무산 ('24~'25年)"),
    ('lv4',"전략적 매각·자본구조 재조정·IPO 재추진 등 검토하였으나 성과 부재"),
    ('lv4',"AI 기반 경쟁 심화: Qualtrics 등 경쟁사의 AI 기능 강화로 고객 이탈 가속"),
    ('lv4',"성장률 둔화: 이자 부담 증가로 현금흐름 악화 → 혁신 투자 축소 → 상품 경쟁력 저하"),
    ('lv3',"3) 연간 이자비용($0.3B)이 EBITDA($0.2B)를 초과 (직전 1년 기준)"),
    ('lv4',"EBITDA만 보면 흑자 구조처럼 보이나, 실제 현금흐름 기준으로는 구조적 부실 상태에 진입"),
    ('lv3',"4) 경영권 채권단 이전 및 투자 손실 확정 ('26.4月)"),
    ('lv4',"Debt-to-Equity Swap: 채권단 보유 $2.8B 채권을 지분 전환하여 채권단이 최대주주가 됨"),
    ('lv4',"PE 투자 역사상 최대 규모 SaaS 차입매수(LBO) 실패 사례"),
    # ── 4. 주요 분석 ──
    ('lv1','4. 주요 분석'),
    ('lv2','실패 원인'),
    ('lv3',"① 성장률 과대 추정: 고성장 지속을 전제로 9배 멀티플(EV/Revenue) 적용"),
    ('lv3',"② 취약한 부채 구조: 변동금리 인상, PIK 누적으로 부채 급증"),
    ('lv3',"③ 이자비용 부담으로 인한 상품 경쟁력 저하: 잉여현금흐름(FCF) 전액을 이자 상환에 투입하여 투자 여력 부족"),
    ('lv3',"④ 투자금 회수 무산: M&A 매각 및 IPO 재진입 실패 — 과도한 부채 구조로 잠재 인수자 협상 난항"),
    ('lv2','예상 파급 효과'),
    ('lv3',"Medallia: 투자 축소로 서비스 경쟁력 저하 우려, 상품 로드맵 불확실, 경쟁사로 이탈 가능성"),
    ('lv3',"사모펀드(PE): 유사한 구조로 인수된 SaaS 기업 리스크 재조명 (별첨 3 참조)"),
    ('lv3',"민간 대출 시장: Blackstone 등 채권자도 원금 회수 불확실성 확대 → 민간 대출 시장 리스크 현실화"),
    ('lv3',"SaaS M&A 시장: 고부채 SaaS 매물 증가 가능성, 전략적 인수자 중심 저가 인수 기회 확대"),
    # ── 5. 시사점 ──
    ('lv1','5. 시사점'),
    ('lv2','SaaS 기업가치 평가시, 성장과 현금창출능력을 함께 고려 필요'),
    ('lv3',"ARR성장률, Rule of 40, NRR 등 SaaS 특화 지표 뿐만 아니라 잉여현금흐름(FCF) 및 부채상환능력(EV/Par Coverage) 지표 병행관리 필요"),
    ('lv2','향후 M&A 검토 시 현금흐름, 부채상환능력 지표도 병행 검토 필요'),
    ('lv2','사업건전성 대시보드 운영 시, 현금흐름 관리 체계 강화 검토 필요'),
    ('lv3',"잉여현금흐름(FCF) 지표 시범 산출 → 사업부 관리 협의 → 대시보드 적용 기준 확정"),
    ('lv3',"Rule of 40·NRR·잉여현금흐름(FCF) 등 통합 모니터링 체계 검토"),
    ('lv3',"(1) FCF = 영업현금흐름 − CAPEX(유무형자산취득액)"),
    ('lv3',"(2) EV/Par Coverage Ratio = EV ÷ Loan Par(누적 차입금·적자금액) → 1 이상 유지 필요"),
    ('lv4',"1 미만 하락 시, 기업가치가 누적 부채를 충분히 커버하지 못하는 구조적 위험 신호"),
    ('lv4',"Medallia: EV 급락 + Loan Par $2.8B 팽창으로 임계치 이하 하락 → 경영권 이전"),
    ('lv5',"내부 적용 기준(안): Annual FCF = 영업현금흐름 / Multiple = 국내 B2B SaaS 시장 평균(5~8배 예상) / Loan Par = 누적 차입금 또는 누적 적자 금액 — 시범 산출 후 사업부 협의 후 확정"),
    # ── 별첨 1 ──
    ('pagebreak',),
    ('lv1','별첨 1. Medallia 상세 프로필'),
    ('lv2','기업 기본 정보'),
    ('table',
     ['항목','내용'],
     [
         ['기업명','Medallia, Inc.'],
         ['설립',"'01年 (미국 캘리포니아)"],
         ['창립자','Borge Hald, Amy Pressman (Fortune 500 기업 임원 컨설팅 경험 기반 창업)'],
         ['본사','San Francisco, CA'],
         ['임직원 수',"약 2,037명 ('21.1月 기준)"],
         ['상장 이력',"'19.7月 NYSE 상장 (MDLA, 공모가 $21) → '21.10月 Thoma Bravo 인수로 상장폐지"],
         ['현재 상태',"채권단(Blackstone 等) 관리하 구조조정 中 ('26年)"],
     ]),
    ('lv2','제품·솔루션 구성'),
    ('table',
     ['카테고리','주요 솔루션'],
     [
         ['고객경험(CX)','고객경험 관리(CEM), 디지털 경험 분석, 경험 오케스트레이션, 개인화 메시징'],
         ['직원경험(EX)','직원 리스닝(Listening), 직원 활성화(Activation) 플랫폼'],
         ['컨택센터','대화형 인텔리전스, 상담원 코칭, 품질 관리, 지능형 콜백'],
         ['AI·분석',"Frontline-Ready AI™ — 별도 교육 없이 현장 즉시 활용 가능한 내장형 AI"],
         ['시장조사','애자일 리서치(Agile Research), 비디오 조사'],
     ]),
    ('lv2','주요 고객 및 시장 지위'),
    ('lv3',"주요 글로벌 고객: Airbnb, Volvo Cars, 3M, Johnson & Johnson, MetLife, PetSmart 等 Fortune 500 다수"),
    ('lv3',"주요 산업군: 금융 서비스, 항공, 호텔·숙박, 통신, 소매, B2B 기업"),
    ('lv3',"Gartner Magic Quadrant '26年 'Leader' 선정 (CX 관리 플랫폼 부문)"),
    ('lv3',"주요 경쟁사: Qualtrics(SAP 계열), NICE Satmetrix, Salesforce Experience Cloud, Sprinklr"),
    ('lv2','재무 현황 (상장폐지 직전 공개 기준)'),
    ('table',
     ['항목','수치','기준'],
     [
         ['연간 매출','$477M',"FY2020 (상장 당시 공개)"],
         ['순손실','-$149M','FY2020'],
         ['인수 후 추정 매출','약 $700M대',"'21年 인수 당시 기준 (추정)"],
     ]),
    ('lv5',"'21年 상장폐지 후 재무 정보 비공개. 이자비용 급증으로 인한 FCF 악화는 기사(OnlyCFO, '26.5月) 분석 기반"),
    # ── 별첨 2 ──
    ('pagebreak',),
    ('lv1','별첨 2. Thoma Bravo 상세 프로필'),
    ('lv2','기업 기본 정보'),
    ('table',
     ['항목','내용'],
     [
         ['기업명','Thoma Bravo, LP'],
         ['설립',"'08年 공식 설립 (전신: '80年 설립된 Golder Thoma & Co.)"],
         ['본사','San Francisco, CA / Chicago, IL'],
         ['운용자산(AUM)',"$183B 이상 ('25.12.31 기준)"],
         ['누적 거래 건수','565건 이상 완료, 누적 대표 가치 $305B 이상'],
         ['특화 분야','엔터프라이즈 소프트웨어, 인프라 소프트웨어, 사이버보안, 기술 서비스'],
     ]),
    ('lv2','주요 포트폴리오'),
    ('table',
     ['기업명','분야','인수금액'],
     [
         ['Proofpoint','엔터프라이즈 사이버보안','$12.3B'],
         ['Anaplan','기업 계획(Planning) 플랫폼','$10.7B'],
         ['Stamps.com','우편·배송 SaaS','$6.6B'],
         ['Medallia','CX 관리 플랫폼','$6.4B (손실 확정)'],
         ['Coupa Software','조달 자동화','$6.2B'],
         ['Sophos','엔드포인트 보안','$3.9B'],
     ]),
    ('lv2','인수 실패 메커니즘 요약'),
    ('lv3',"인수 時 LIBOR 0.75% 변동금리로 $1.8B 조달 → 연준 금리 인상으로 SOFR 5.4% 도달"),
    ('lv3',"연이자 $135M → $300M 급등 → PIK 이자 누적으로 원금 $1.8B → $2.2B (추가 차입 포함 $2.8B)"),
    ('lv3',"EBITDA $200M으로 이자($300M)조차 충당 불가 → 채무불이행(Default) → 경영권 상실"),
    # ── 별첨 3 ──
    ('pagebreak',),
    ('lv1',"별첨 3. 동일 구조 위험 노출 대표 사례 ('21~'22年)"),
    ('table',
     ['기업명','PE 인수사','인수 시기','인수금액','Multiple(Revenue 기준)','사업 영역'],
     [
         ['Proofpoint','Thoma Bravo',"'21.4月",'$12.3B','10.0~11.0','사이버보안 SaaS'],
         ['RealPage','Thoma Bravo',"'21.4月",'$10.2B','8.5~9.0','부동산 관리 SaaS'],
         ['Zendesk','Permira·H&F',"'22.11月",'$10.2B','6.5~7.0','CX SaaS'],
         ['Avalara','Vista Equity',"'22.10月",'$8.4B','9.0~10.0','세금 자동화 SaaS'],
         ['SailPoint','Thoma Bravo',"'22.4月",'$6.9B','11.0~12.0','ID 보안 SaaS'],
         ['Stamps.com','Thoma Bravo',"'21.10月",'$6.6B','8.5~9.0','우편·배송 SaaS'],
     ]),
    # ── 별첨 4 ──
    ('lv1','별첨 4. 참고문헌'),
    ('lv2','분석 기사'),
    ('lv3',"OnlyCFO, \"Thoma Bravo's $5B Medallia Loss | The PE Reckoning is Coming...\" ('26.5月) — onlycfo.io"),
    ('lv3',"Bloomberg, \"Thoma Bravo Declines to Invest More in Troubled Medallia\" ('26.4月) — bloomberg.com"),
    ('lv2','기업 공식 자료'),
    ('lv3',"Medallia 공식 홈페이지 — medallia.com"),
    ('lv3',"Thoma Bravo 공식 홈페이지 — thomabravo.com"),
    ('lv3',"Thoma Bravo, \"Medallia to be Acquired by Thoma Bravo for $6.4 Billion\" (보도자료, '21.7月)"),
    ('lv3',"PR Newswire, \"Medallia Announces Pricing of Initial Public Offering\" ('19.7月)"),
    ('lv2','개념 참고'),
    ('lv3',"Wikipedia, \"Medallia\" — en.wikipedia.org"),
    ('lv3',"Wikipedia, \"Thoma Bravo\" — en.wikipedia.org"),
    ('lv3',"Wikipedia, \"Rule of 40\" — en.wikipedia.org"),
    ('lv3',"McKinsey, \"SaaS and the Rule of 40\" — mckinsey.com"),
]

# ── 문서 빌드 ─────────────────────────────────────────────────────────────────

def build(output_path):
    doc = setup_doc()
    add_footer(doc)

    # 제목
    p = doc.add_paragraph()
    set_pf(p, before=0, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, 'Medallia 사례로 본 SaaS 기업 가치 평가 기준',
            size=PT_TITLE, bold=True, underline=True)

    # 날짜 · 소속
    p = doc.add_paragraph()
    set_pf(p, before=0, after=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p, "'26.05.14  솔루션전략그룹(솔루션)")

    # 콘텐츠 렌더
    for item in REPORT:
        k = item[0]
        if k == 'summary':
            add_summary_table(doc, item[1])
        elif k == 'lv1':
            p = doc.add_paragraph(); set_pf(p, before=16, after=12)
            add_run(p, item[1], bold=True)
        elif k == 'lv2':
            p = doc.add_paragraph(); set_pf(p, before=0, after=9)
            add_run(p, ' □ ' + item[1], bold=True)
        elif k == 'lv3':
            p = doc.add_paragraph(); set_pf(p, before=0, after=9)
            add_run(p, '   - ' + item[1])
        elif k == 'lv4':
            p = doc.add_paragraph(); set_pf(p, before=0, after=9)
            add_run(p, '     · ' + item[1])
        elif k == 'lv5':
            p = doc.add_paragraph(); set_pf(p, before=0, after=9)
            add_run(p, '   ※ ' + item[1])
        elif k == 'arrow':
            p = doc.add_paragraph(); set_pf(p, before=0, after=9)
            add_run(p, '→ ' + item[1], bold=True)
        elif k == 'table':
            add_table(doc, item[1], item[2])
            gap = doc.add_paragraph(); set_pf(gap, after=6)
        elif k == 'pagebreak':
            page_break(doc)

    # 끝 표시
    p = doc.add_paragraph()
    set_pf(p, before=12, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p, '- 以 上 -')

    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    OUT = '/home/ubuntu/files/medallia_report_final.docx'
    try:
        result = build(OUT)
        size = os.path.getsize(result)
        print(f'완료: {result}  ({size:,} bytes, {size/1024:.1f} KB)')
    except Exception:
        import traceback; traceback.print_exc(); sys.exit(1)
