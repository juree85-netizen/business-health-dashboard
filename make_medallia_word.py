#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Medallia 보고서 → Word 변환 스크립트 (python-docx)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 설정 (A4, 여백) ──────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.5)

# ── 기본 스타일 ──────────────────────────────────────────
def set_font(run, size=12, bold=False, color=None, name='바탕'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
             left_indent=0, line_spacing=1.5):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.left_indent  = Cm(left_indent)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

def add_title(text, size=18):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
    r = p.add_run(text)
    set_font(r, size=size, bold=True)
    r.font.underline = True
    return p

def add_ch(text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=10, space_after=4)
    r = p.add_run(text)
    set_font(r, size=13, bold=True)
    return p

def add_sq(text):
    p = doc.add_paragraph()
    para_fmt(p, left_indent=0.5, space_before=5, space_after=2)
    r = p.add_run('□ ' + text)
    set_font(r, size=12)
    return p

def add_dash(text, indent=1.5):
    p = doc.add_paragraph()
    para_fmt(p, left_indent=indent, space_before=0, space_after=2)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(-0.5)
    r = p.add_run('- ' + text)
    set_font(r, size=11.5)
    return p

def add_dot(text, indent=2.5):
    p = doc.add_paragraph()
    para_fmt(p, left_indent=indent, space_before=0, space_after=2)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(-0.5)
    r = p.add_run('· ' + text)
    set_font(r, size=11)
    return p

def add_fn(text):
    p = doc.add_paragraph()
    para_fmt(p, left_indent=2.0, space_before=1, space_after=2)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(-0.7)
    r = p.add_run('※ ' + text)
    set_font(r, size=9.5, color=(80,80,80))
    return p

def add_body(text, indent=1.5):
    p = doc.add_paragraph()
    para_fmt(p, left_indent=indent, space_before=1, space_after=3)
    r = p.add_run(text)
    set_font(r, size=11.5)
    return p

def add_sub_heading(text, indent=1.5):
    p = doc.add_paragraph()
    para_fmt(p, left_indent=indent, space_before=5, space_after=2)
    r = p.add_run(text)
    set_font(r, size=12, bold=True)
    return p

def add_sep():
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=4)
    p.add_run('─' * 55)
    return p

def add_page_break():
    doc.add_page_break()

# 테이블 헬퍼
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  color=None, font_color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=font_color)
    if color:
        set_cell_bg(cell, color)
    return cell

def add_tbl(headers, rows, col_widths=None, hdr_color='1e1e1e', alt_color=None):
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_text(c, h, size=10.5, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER,
                      color=hdr_color, font_color=(255,255,255))

    # 데이터
    for ri, row in enumerate(rows):
        bg = alt_color if (alt_color and ri % 2 == 1) else None
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            txt = str(val) if val is not None else ''
            align = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            fc = (176,0,0) if txt.startswith('-') and any(ch.isdigit() for ch in txt) else None
            if txt.endswith('×') and txt[0] == '0':
                fc = (176,0,0)
            set_cell_text(c, txt, size=10.5, align=align, color=bg, font_color=fc)

    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

def add_app_title(text):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=8)
    r = p.add_run(text)
    set_font(r, size=14, bold=True)
    return p

def add_app_sub(text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=8, space_after=3)
    r = p.add_run(text)
    set_font(r, size=12, bold=True)
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ════════════════════════════════════════════════════════
# 목차
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_fmt(p, space_before=0, space_after=6)
r = p.add_run('목    차')
set_font(r, size=13, bold=True)

toc_items = [
    ('1장', '보고 개요', '분석 배경 / 분석 범위 / 분석 목적'),
    ('2장', 'Medallia 개요', '기업 개요 / 재무 현황 (인수 전)'),
    ('3장', 'Medallia 투자/인수 현황', "Thoma Bravo의 Medallia 인수 ('21년) / 인수 이후 상황"),
    ('4장', '예상 파급 효과', ''),
    ('5장', '시사점', ''),
]
for num, title, sub in toc_items:
    p = doc.add_paragraph()
    para_fmt(p, space_before=3, space_after=1)
    r1 = p.add_run(f'{num}  ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(title)
    set_font(r2, size=12, bold=True)
    if sub:
        p2 = doc.add_paragraph()
        para_fmt(p2, left_indent=1.2, space_before=0, space_after=2)
        r = p2.add_run(sub)
        set_font(r, size=10.5, color=(80,80,80))

p = doc.add_paragraph()
para_fmt(p, space_before=8, space_after=3)
r = p.add_run('별   첨')
set_font(r, size=12, bold=True)

app_items = [
    '[별첨 1]  Medallia 상세 프로필',
    '[별첨 2]  Thoma Bravo 상세 프로필',
    "[별첨 3]  LBO(차입매수)를 통한 인수 사례 ('21~'22년)",
    '[별첨 4]  이해관계자 관점 분석 : Thoma Bravo vs Blackstone',
    '[별첨 5]  Qualtrics vs Medallia 비교',
    "[별첨 6]  SaaS 기업의 M&A 검토 시 평가 기준(안)",
    '[별첨 7]  PIK 원금 누적 시뮬레이션 및 EV/Par Coverage 이탈 과정',
    '[별첨 8]  SaaS 기업가치 평가의 새로운 기준 : Gen AI 대체 가능성',
    '참고문헌',
]
for item in app_items:
    p = doc.add_paragraph()
    para_fmt(p, left_indent=1.2, space_before=1, space_after=1)
    r = p.add_run(item)
    set_font(r, size=10.5)

add_page_break()


# ════════════════════════════════════════════════════════
# PAGE 1  1장 + 2장
# ════════════════════════════════════════════════════════
add_title('Medallia 사례로 본 SaaS 기업 가치 평가')
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=10)
r = p.add_run("'26.05.19    솔루션전략그룹(솔루션)")
set_font(r, size=11)

add_ch('1. 보고 개요')
add_sq('분석 배경')
add_dash("Thoma Bravo는 Medallia를 $6.4B(EV/Revenue 멀티플 약 9배)에 인수하며 상장폐지 ('21.7월)")
add_dash("美 연준 금리 인상으로 SOFR가 약 5.4%까지 상승, 차입 이자 부담 급증 ('22~'25년)")
add_dash("연간 이자비용($0.3B)이 EBITDA($0.2B)를 초과하는 구조적 부실 직면, 채권단으로 경영권 이전 ('26.4월)")

add_sq('분석 범위')
add_dash('분석 대상: Medallia / Thoma Bravo')
add_dash("분석 기간: '21년 인수 ~ '26년 채권단 경영권 이전")
add_dash('재무 수치 출처: Bloomberg, OnlyCFO 분석 기반 추정치')

add_sq('분석 목적')
add_dash('PE 기반 SaaS 차입매수(LBO) 방식의 실패 원인 분석')
add_dash('SaaS 기업가치 및 M&A 추진 시 평가 기준 재점검')

add_sep()

add_ch('2. Medallia 개요')
add_sq('기업 개요')
add_dash("'01년 설립, 샌프란시스코 소재 CX(고객경험)관리 SaaS 기업, VoC 플랫폼 부문 글로벌 2위 (1위 Qualtrics)")
add_dot('CRM 시장은 Sales, Marketing, Customer Service & Support, Digital Commerce 4개 세그먼트로 구분, Medallia는 Customer Service & Support 중 VoC 플랫폼 부문 (출처: Gartner Magic Quadrant 2025·2026)')
add_dash('오퍼링 : 고객경험 피드백 수집·분석, NPS 측정, 직원경험 플랫폼, 옴니채널 고객 인텔리전스')
add_dash("주요 고객 : 금융·항공·호텔·통신 등 약 1,000개 고객사 보유")
add_dash("'15~'18년 : VC, PE로부터 복수의 투자 유치, 기업가치 급성장")
add_dash("'19년 NYSE 상장 (IPO 공모가 $21), 상장 이후 $16~$48 수준")
add_dash("'21년 Thoma Bravo에 인수되며 상장폐지")

add_sq('재무 현황 (인수 전)')
add_tbl(
    headers=['항목', '수치', '시점', '비고'],
    rows=[
        ['매출', '$477M', 'FY2021', '23~26%(YoY) 고성장'],
        ['EBITDA', '-$109M', 'FY2021', ''],
        ['FCF', '-$19M', 'FY2021', 'CAPEX, 운전자본, 현금이자 반영 시 적자'],
        ['Rule of 40', '10', "'21년", 'SaaS기업 중 하위 10% 수준 (기준: 40 이상)'],
        ['NRR', '116~126%', "'18~'19년", "'18.1월 126%→'19.1월 116%→'19.4월 119%\n상장 전 하락세, 인수 이후 추가 하락"],
    ],
    col_widths=[2.8, 2.0, 2.2, 8.5]
)
add_fn('Free Cash Flow(잉여현금흐름) = OCF(영업활동현금흐름) - CapEx(설비투자)')
add_fn('Net Revenue Retention : 순매출유지율 (기존고객의 현재매출 / 기존고객의 과거매출) × 100')

add_page_break()


# ════════════════════════════════════════════════════════
# PAGE 2  3장 — TB 인수
# ════════════════════════════════════════════════════════
add_ch('3. Medallia 투자/인수 현황')
add_sq("Thoma Bravo의 Medallia 인수 ('21년)")
add_dash('당시 고성장 SaaS 프리미엄을 고려, EV/Revenue 멀티플을 약 9배로 적용')
add_fn('멀티플 : 기업가치(Enterprise Value)를 특정 재무지표(매출/이익/현금흐름)로 나눈 값으로, 시장이 해당 기업 실적에 부여하는 가치 수준을 의미')

# 자금 조달 구조 테이블
add_sub_heading('▸ 자금 조달 구조')
tbl = doc.add_table(rows=2, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
# 1행: 총 인수가격 병합
row0 = tbl.rows[0]
row0.cells[0].merge(row0.cells[2])
set_cell_text(row0.cells[0], '총 인수가격  $6.4B', size=11.5, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, color='f0f0f0')
# 2행: equity + + debt
set_cell_text(tbl.rows[1].cells[0], '투자금 (Equity)\n$4.6B\nThoma Bravo 등',
              size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color='eef3fb')
set_cell_text(tbl.rows[1].cells[1], '+', size=14, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(tbl.rows[1].cells[2], '변동금리 차입금 (Debt)\n$1.8B\n민간대출업자 (Blackstone 등)',
              size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color='fef2f2')
for row in tbl.rows:
    row.cells[0].width = Cm(6.5)
    row.cells[1].width = Cm(1.0)
    row.cells[2].width = Cm(6.5)

add_fn("당시 LIBOR 0.75%. LIBOR : 런던 은행 간 단기 대출 금리, '23년 공식 폐지 후 SOFR로 대체됨.")
add_dash("인수 당시 계획 : 저금리 환경에서 낮은 비용으로 차입 매수, 영업현금흐름으로 이자 감당, 3~5년 내 매각 또는 IPO 재상장으로 투자금 회수")

add_sub_heading('▸ 인수 판단 기준')
add_dot('Medallia 인수 당시의 9배 멀티플은 상대적으로 저렴한 투자로 평가')
add_dot('Medallia는 ARR 성장률 23~26%의 고성장 기업으로 FCF는 적자였으나 향후 개선 전망')
add_fn('단, Medallia의 Rule of 40, FCF 마진은 동종업계 하위 10% 수준임을 인지')
add_dot('당시 낮은 이자 부담으로, PIK를 활용한 차입매수(LBO) 기반 인수 결정')

add_tbl(
    headers=['연도', '매출', 'EBITDA', 'FCF', 'FCF Margin'],
    rows=[
        ["FY2019 ('18년)", '$314M', '-$67M', '-$27M', '-8.6%'],
        ["FY2020 ('19년)", '$403M', '-$99M', '-$24M', '-6.0%'],
        ["FY2021 ('20년)", '$477M', '-$109M', '-$19M', '-4.0%'],
        ["FY2022 ('21년)", '$560M', '-$125M', '-$115M', '-20.4%'],
    ],
    col_widths=[3.5, 2.2, 2.2, 2.2, 2.4]
)
add_fn("매출성장률 ('18~'21년) 21.2%로 성장하였으나 FCF Margin은 -11.8%p로 역성장, '21년은 인수 후 이자부담으로 FCF 악화")

add_page_break()


# ════════════════════════════════════════════════════════
# PAGE 3  3장 — 인수 이후 상황 1) + 2)
# ════════════════════════════════════════════════════════
add_sq('인수 이후 상황')
add_sub_heading("1) 금리 인상 및 취약한 부채 구조로 이자부담 급증 ('22년 이후)")
add_dash('이자 급증 : SOFR 연동 변동금리 영향으로 이자비용 급등')
add_fn('SOFR : 미국 달러 단기 금리 기준, LIBOR를 대체하며 미 연준 금리 결정에 연동')

add_tbl(
    headers=['구분', "'21년", "'23년", "'26년(예상)"],
    rows=[
        ['기준금리(SOFR)', '0.75%', '5.4%', '3.7%'],
        ['연이자', '$135M', '$300M', '$270M'],
    ],
    col_widths=[4.0, 3.0, 3.0, 3.0]
)
add_dash('부채 규모 확대 : PIK 적용으로 원금 $1.8B → 약 $2.2B (현금 미지급 이자가 원금에 누적 반영)')
add_fn("PIK 비중 점진 축소 ('21~'22년 100% → '22~'23년 50% → '23~'25년 35% → '25~'26년 0%)")
add_fn("Medallia를 통한 추가 M&A 관련 차입금($0.6B) 반영 시 총부채 약 $2.8B 수준")

add_sub_heading('2) 현금흐름 악화 (FCF Coverage 악화)')
add_tbl(
    headers=['연도', '매출(추정)', 'EBITDA(추정)', 'FCF(추정)', '연이자', 'FCF Coverage', '비고'],
    rows=[
        ["FY2022 ('21년)", '$560M', '-$125M', '-$115M', '$135M', '0.85×', '인수'],
        ["FY2023 ('22년)", '$685M', '$25M', '-$10M', '$180M', '0.06×', '구조적 부실 진입'],
        ["FY2024 ('23년)", '$740M', '$155M', '-$10M', '$300M', '0.03×', ''],
        ["FY2025 ('24년)", '$740M', '$200M', '-$5M', '$280M', '0.02×', '비용 절감 효과'],
        ["FY2026 ('25년)", '$735M', '$200M', '-$75M', '$270M', '0.28×', '경영권 이전 결정'],
    ],
    col_widths=[3.2, 1.9, 2.0, 1.8, 1.7, 2.0, 2.4]
)
add_dash("매출은 '21~'25년 CAGR 7.0%로 성장, FCF 개선 중이었으나 연이자 2배 이상 증가로 현금흐름 악화")
add_dash("'25년 연간 이자비용($270M)이 EBITDA($200M)를 초과")
add_dash("인수 당시부터 FCF Coverage 1.0× 미만 — 인수 첫 해('21년) FCF $115M, 이자 $135M → Coverage 0.85×")
add_fn('FCF Coverage = FCF ÷ 연이자. 1.0× 미만 = 이자를 FCF로 감당 불가')

add_page_break()


# ════════════════════════════════════════════════════════
# PAGE 4  3장 — 인수 이후 상황 3) + 4)
# ════════════════════════════════════════════════════════
add_sub_heading("3) 매각, IPO 재진입 실패로 투자금 회수 무산 ('24~'25년)")
add_dash('전략적 매각 불가 : 매각가 $5B 산정하였으나 Gen AI 위협·고부채 구조로 전략적 인수자 부재')
add_dash('IPO 재진입 불가 : 불안정 FCF, 성장률 둔화, SaaS 멀티플 수축기로 투자자 확보 어려움')
add_dash('자본구조 재조정 실패 : 추가 자본 조달 불가, 채권단(Blackstone) 부채 감면 거부')
add_dash('상품 경쟁력 약화 : Qualtrics 등 경쟁사의 AI 기능 강화로 고객 이탈 가속')
add_dash('성장률 둔화 : 이자 부담 증가 → 현금흐름 악화 → 혁신 투자 축소 → 상품 경쟁력 저하')

add_sub_heading("4) 경영권 채권단 이전 및 투자 손실 확정 ('26.4월)")
add_dash('경영권 이전 배경')
add_dot('채권단의 현금 이자 지급 압박 (PIK 유예 중단)')
add_dot('인수 당시 고평가 인정과 성장의 한계')
add_dot('생성형 AI에 의해 대체 가능성이 높은 SaaS 기업 (SaaSpocalypse)')
add_fn('SaaSpocalypse : 고금리, AI대체 가능성, 성장둔화, 멀티플 붕괴로 인해 SaaS 기업들의 가치와 생존 환경이 급격히 악화된 상황')
add_dot('대체 투자 기회의 수익률을 고려하여 기존 투입자금은 손실 처리')
add_dash('채권단 보유 $2.8B 채권을 지분 전환 → 채권단이 최대주주가 됨 (Debt-to-Equity Swap)')
add_dash('PE 투자 역사상 최대 규모 SaaS 차입매수(LBO) 실패 사례')

add_page_break()


# ════════════════════════════════════════════════════════
# PAGE 5  4장 예상 파급 효과
# ════════════════════════════════════════════════════════
add_ch('4. 예상 파급 효과')
add_sq('Medallia')
add_body('투자 축소로 서비스 경쟁력 저하 우려, 상품 로드맵 불확실, Medallia 고객의 Qualtrics·Salesforce 등 경쟁사로 이탈')

add_sq('사모펀드(PE)')
add_body('유사한 구조로 인수된 SaaS 기업 리스크 재조명')
add_dot('Zendesk, Proofpoint 등 유사 사례 확산 가능성 (별첨3 참조)')

add_sq('민간 대출 시장')
add_body('Blackstone 등 채권자도 원금 회수 불확실성 확대 → 민간 대출 시장 리스크 현실화')
add_dot('① 금리 상승 시 이자수익 증가로 설계하였으나 Medallia의 부채 상환 능력을 떨어뜨리는 결과 초래 ($1.8B → $2.8B)')
add_dot('② 부채 감면 거부 후 경영권을 인수하였으나 매각가치가 $2.8B 미만으로 추정되어 원금 손실 가능성 높음')

add_sq('SaaS M&A 시장')
add_body('SaaS의 밸류에이션이 조정되는 과정으로 고부채 SaaS매물 증가 가능성')

add_page_break()


# ════════════════════════════════════════════════════════
# PAGE 6  5장 시사점
# ════════════════════════════════════════════════════════
add_ch('5. 시사점')
add_sq('SaaS기업의 M&A 검토 시 현금창출능력 고려 必')
add_dash('현금창출능력을 평가하기 위해서는 잉여현금흐름(FCF), 부채상환능력(EV/Par Coverage) 및 Rule of 40도 함께 평가')
add_dot('잉여현금흐름(FCF) : 이자상환능력(FCF Coverage) 최소 1 이상')
add_dot('성장+수익 균형 : Rule of 40 최소 30 이상')
add_dot('부채상환능력(EV/Par Coverage) : 최소 2 이상')
add_dot('고객유지력 : NRR 최소 110% 이상')

add_sq('솔루션의 Gen AI대체 가능성 판단 기준 수립 必')
add_dash('데이터 독점/폐쇄성 : 특정 기업 보유 데이터 혹은 계약/규제로 접근이 어려운 데이터일 경우 대체 가능성 低')
add_dash('사용량 또는 성과 기반 가격 체계 : 대체 가능성 低')
add_dash('치명적 리스크가 있는 분야 : 고도의 전문지식, 높은 정확성/신뢰성 요구 분야 대체가능성 低')

add_sq('자사솔루션에 대한 현금흐름 모니터링 체계 필요')
add_dash('매출/영업이익 등 경영실적 외에도 Rule of 40, 잉여현금흐름(FCF), NRR 등 통합 모니터링 (사업건전성 대시보드 활용)')

p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=20, space_after=0)
r = p.add_run('- 以 上 -')
set_font(r, size=12, bold=True)

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 1 — Medallia 상세 프로필
# ════════════════════════════════════════════════════════
add_app_title('【 별첨 1 】 Medallia 상세 프로필')

add_app_sub('기업 기본 정보')
add_tbl(
    headers=['항목', '내용'],
    rows=[
        ['기업명', 'Medallia, Inc.'],
        ['설립', "'01년 (미국 캘리포니아)"],
        ['창립자', 'Borge Hald, Amy Pressman (Fortune 500 기업 임원 컨설팅 경험 기반 창업)'],
        ['본사', 'San Francisco, CA'],
        ['임직원 수', "약 2,037명 ('21.1월 기준)"],
        ['상장 이력', "'19.7월 NYSE 상장 (MDLA, 공모가 $21) → '21.10월 Thoma Bravo 인수로 상장폐지"],
        ['현재 상태', "채권단(Blackstone 等) 관리하 구조조정 中 ('26년)"],
    ],
    col_widths=[3.0, 12.5]
)

add_app_sub('제품·솔루션 구성')
add_tbl(
    headers=['카테고리', '주요 솔루션'],
    rows=[
        ['고객경험 (CX)', '고객경험 관리(CEM), 디지털 경험 분석, 경험 오케스트레이션, 개인화 메시징'],
        ['직원경험 (EX)', '직원 리스닝(Listening), 직원 활성화(Activation) 플랫폼'],
        ['컨택센터', '대화형 인텔리전스, 상담원 코칭, 품질 관리, 지능형 콜백'],
        ['AI·분석', 'Frontline-Ready AI™ — 별도 교육 없이 현장 즉시 활용 가능한 내장형 AI'],
        ['시장조사', '애자일 리서치(Agile Research), 비디오 조사'],
    ],
    col_widths=[3.5, 12.0]
)

add_app_sub('주요 고객 및 시장 지위')
add_dash('주요 글로벌 고객 : Airbnb, Volvo Cars, 3M, Johnson & Johnson, MetLife, PetSmart 等 Fortune 500 다수')
add_dash('주요 산업군 : 금융 서비스, 항공, 호텔·숙박, 통신, 소매, B2B 기업')
add_dash("Gartner Magic Quadrant '26년 'Leader' 선정 (CX 관리 플랫폼 부문)")
add_dash('주요 경쟁사 : Qualtrics(SAP 계열), NICE Satmetrix, Salesforce Experience Cloud, Sprinklr')

add_app_sub('재무 현황 (상장폐지 직전 공개 기준)')
add_tbl(
    headers=['항목', '수치', '기준'],
    rows=[
        ['연간 매출', '$477M', 'FY2020 (상장 당시 공개)'],
        ['순손실', '-$149M', 'FY2020'],
        ['인수 후 추정 매출', '약 $700M대', "'21년 인수 당시 기준 (추정)"],
    ],
    col_widths=[4.0, 3.0, 8.5]
)
add_fn("'21년 상장폐지 후 재무 정보 비공개. 이자비용 급증으로 인한 FCF 악화는 기사 분석 기반")

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 2 — Thoma Bravo 상세 프로필
# ════════════════════════════════════════════════════════
add_app_title('【 별첨 2 】 Thoma Bravo 상세 프로필')

add_app_sub('기업 기본 정보')
add_tbl(
    headers=['항목', '내용'],
    rows=[
        ['기업명', 'Thoma Bravo, LP'],
        ['설립', "'08년 공식 설립 (전신: '80년 설립된 Golder Thoma & Co.)"],
        ['본사', 'San Francisco, CA / Chicago, IL'],
        ['운용자산(AUM)', "$183B 이상 ('25.12.31 기준)"],
        ['누적 거래 건수', '565건 이상 완료, 누적 대표 가치 $305B 이상'],
        ['특화 분야', '엔터프라이즈 소프트웨어, 인프라 소프트웨어, 사이버보안, 기술 서비스'],
    ],
    col_widths=[3.5, 12.0]
)

add_app_sub('주요 포트폴리오')
add_tbl(
    headers=['기업명', '분야', '인수금액'],
    rows=[
        ['Proofpoint', '엔터프라이즈 사이버보안', '$12.3B'],
        ['Anaplan', '기업 계획(Planning) 플랫폼', '$10.7B'],
        ['Stamps.com', '우편·배송 SaaS', '$6.6B'],
        ['Medallia', 'CX 관리 플랫폼', '(손실 확정) $6.4B'],
        ['Coupa Software', '조달 자동화', '$6.2B'],
        ['Sophos', '엔드포인트 보안', '$3.9B'],
    ],
    col_widths=[4.0, 7.0, 4.5]
)

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 3 — LBO 인수 사례
# ════════════════════════════════════════════════════════
add_app_title("【 별첨 3 】 LBO(차입매수)를 통한 인수 사례 ('21~'22년)")

add_tbl(
    headers=['기업명', 'PE 인수사', '인수시기', '인수금액', 'Multiple\n(Revenue)', 'FCF\nCoverage', '사업 영역', "현재 상태('26년)"],
    rows=[
        ['Proofpoint', 'Thoma Bravo', "'21.4", '$12.3B', '10~11×', '1.0×', '사이버보안 SaaS', '재상장 검토 중. 추정 부채 $5.3B. 사이버보안 수요 견조'],
        ['RealPage', 'Thoma Bravo', "'21.4", '$10.2B', '8.5~9×', '0.8×', '부동산 관리 SaaS', 'FTC 반독점 소송 진행 중. 소송 리스크 부담'],
        ['Zendesk', 'Permira·H&F', "'22.11", '$10.2B', '6.5~7×', '0.7×', 'CX SaaS', 'AI 전환 진행 중. GenAI 대체 압력 존재. 구조조정 일부 시행'],
        ['Qualtrics', 'Silver Lake·CPP', "'23.6", '$12.5B', '8~9×', '1.0×', 'XM(경험관리) SaaS', 'SAP 분리 후 인수 완료. GenAI 위험 노출. Exit 경로 불확실'],
        ['Avalara', 'Vista Equity', "'22.10", '$8.4B', '9~10×', '1.2×', '세금 자동화 SaaS', '세금 규정 복잡성으로 GenAI 대체 낮음. 비교적 안정적'],
        ['Medallia', 'Thoma Bravo', "'21.7", '$6.4B', '9×', '0.67×→0.4×', 'CX 관리 SaaS', '채권단 경영권 이전 완료. 손실 확정. PE 역사상 최대 SaaS LBO 손실'],
        ['SailPoint', 'Thoma Bravo', "'22.4", '$6.9B', '11~12×', '1.3×', 'ID 보안 SaaS', "재상장 성공('24.5월, NYSE). FCF 개선. 동일 PE의 성공 사례"],
    ],
    col_widths=[1.9, 2.0, 1.3, 1.7, 1.6, 1.6, 2.2, 3.2]
)
add_fn('FCF Coverage = FCF ÷ 연이자 (1.0 미만일 경우 이자를 FCF로 감당 불가)')
add_fn('Medallia 수치만 분석 기사(OnlyCFO) 기반 확인치; 나머지는 공개 재무 데이터 추정치.')

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 4 — 이해관계자 관점 분석
# ════════════════════════════════════════════════════════
add_app_title('【 별첨 4 】 이해관계자 관점 분석 : Thoma Bravo vs Blackstone')

add_app_sub('Thoma Bravo : PE 투자자 (Equity) 관점')
add_tbl(
    headers=['단계', '당시 판단', '사후 평가'],
    rows=[
        ['투자 논거', 'EV/Revenue 9× = 업계 평균(15~20×) 대비 저렴. ARR 23~26% 고성장, Gartner MQ Leader, Fortune 500 고객 기반', '저렴한 멀티플이 낮은 FCF 창출력을 가렸음. FCF Coverage 0.67×는 "위험 신호"'],
        ['FCF 개선 기대', 'TB 플레이북(S&M 절감, 헤드카운트 최적화) 적용 → 3~5년 내 FCF Coverage 1.5× 달성 전망', '실제 FCF 개선 속도 < 이자 상승 속도. 성장 둔화 + 금리 인상이 겹쳐 플레이북 작동 불가'],
        ['금리 리스크', '초저금리(LIBOR 0.75%) 환경에서 연이자 $135M은 감당 가능. 금리 헤지 불필요로 판단', 'SOFR 100% 변동금리 노출 → 이자 $135M→$300M. 금리 스왑 미실행이 결정적 패착'],
        ['Exit 전략', '"3~5년 내 IPO 또는 전략 매각", 시장 회복 시 멀티플 재팽창 기대', 'SaaS 멀티플 수축 + GenAI 대체 압력 + 부채 $2.78B → 4가지 회수 경로 모두 차단'],
        ['추가 투자 거절', "Bloomberg '26.4월: 추가 $500M 투입해도 EV/Par Coverage 1.0× 미달. LP 신뢰 손상 우려 → 거절", '이 결정이 실질적 손실 확정 선언이었음'],
    ],
    col_widths=[2.2, 6.5, 6.8]
)

add_app_sub('Blackstone : 채권자 (Debt) 관점')
add_tbl(
    headers=['단계', '당시 판단', '사후 평가'],
    rows=[
        ['대출 논거', 'TB 포트폴리오사 → PE 최상위 보증 효과. Senior Secured 구조로 원금 보호. 담보: Medallia 전체 자산', 'PE 보증 효과는 TB의 추가 투자 거절로 소멸. EV가 원금 아래 하락 시 담보 효과도 소멸'],
        ['변동금리 설계', 'SOFR 연동: 금리 상승 시 이자 수익 자동 증가 — 대출자에게 구조적으로 유리', '금리 상승으로 이자 수익은 증가했으나 Medallia 상환 능력 파괴 — 자기 파괴적 역설'],
        ['PIK 허용', '초기 현금 부담 완화 → 포트폴리오 생존 기간 연장 → 대출 상환 가능성 제고', 'PIK 복리 누적으로 원금 $1.8B → $2.78B. EV/Par Coverage 급락의 직접 원인'],
        ['D/E Swap 결과', '헤어컷(부채 감면) 거부 → Debt-to-Equity Swap: 명목상 Medallia 대주주 취득', '실제 Exit 가치는 원금 $2.78B 미만으로 추정 — "원하지 않은 Equity"를 보유하게 된 아이러니'],
    ],
    col_widths=[2.2, 6.5, 6.8]
)

add_app_sub('이해관계자별 개선 기준 : 사례 역산')
add_tbl(
    headers=['항목', 'Thoma Bravo (PE 투자자)', 'Blackstone (채권자)'],
    rows=[
        ['FCF Coverage 기준', 'FCF Coverage ≥ 1.5× 이상만 변동금리 LBO 진행', '차입자 FCF Coverage ≥ 1.5× 이상 대출 승인 (미달 시 고정금리 요구)'],
        ['금리 헤지', '변동금리 차입액의 30~50% 금리 스왑(IRS) 헤지 의무화', '차입자에게 금리 캡(Interest Rate Cap) 취득 요구 조항 삽입'],
        ['PIK 제한', 'PIK 비중 50% 이하, 허용 기간 최대 2년으로 계약 제한', 'PIK 허용 기간 최대 2년. 이후 현금 이자 전환 의무화'],
        ['EV/Par 경보', 'EV/Par Coverage 2.0× 이하 → 추가 Equity 투입 또는 Exit 가속화', 'EV/Par Coverage 1.5× 이하 → 조기 상환 요구 또는 금리 재협상 트리거'],
        ['Exit 조건 명시', 'FCF Coverage 2.0× 달성 시 Exit 추진으로 계약에 명시', '대출 만기 내 Exit 미실현 시 조기 상환권 확보'],
        ['GenAI 대체 평가', '핵심 기능 중 LLM 대체 가능 비중 ≥ 40% → 투자 제한 또는 멀티플 할인', 'GenAI 고대체 위험 기업 대출 시 리스크 프리미엄(spread) 추가 100bp 적용'],
    ],
    col_widths=[2.8, 6.6, 6.1]
)

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 5 — Qualtrics vs Medallia
# ════════════════════════════════════════════════════════
add_app_title('【 별첨 5 】 Qualtrics vs Medallia 비교')

add_tbl(
    headers=['항목', 'Qualtrics', 'Medallia'],
    rows=[
        ['설립 / 소재', "'02년 / 미국 Provo", "'01년 / 미국 San Francisco"],
        ['현재 소유구조', "Silver Lake·CPP 컨소시엄 / '23년 SAP에서 분리 인수 ($12.5B)", "채권단(Blackstone 등) 경영권 이전 / '26.4월 Debt-to-Equity Swap 완료"],
        ['연매출 추정', '$1.5B (FY2023 추정)', '$760M (FY2025 추정)'],
        ['인수 Multiple', "EV/Revenue ≈ 8~9× ('23년)", "EV/Revenue ≈ 9× ('21년)"],
        ['핵심 강점', 'XM 플랫폼 통합성 (EX·CX·BX·PX 4축) / SAP ERP 연동 고객 기반', '옴니채널 데이터 파이프라인 / Fortune 500 대기업 워크플로우 내재화'],
        ['상품 구성', 'XM Platform, CoreXM, EmployeeXM, CustomerXM, BrandXM, DesignXM', 'CX·EX 피드백 플랫폼, NPS 측정, Frontline-Ready AI™, 옴니채널 인텔리전스'],
        ['주요 고객층', '엔터프라이즈·중견기업 혼합 (SAP 고객 연계로 다양한 산업)', 'Fortune 500 대기업 집중 (금융·항공·호텔·통신)'],
        ['AI / GenAI 대응', 'AI 기능 강화 지속 투자 (SAP AI Core 연계) / Qualtrics AI™ 출시', 'Frontline-Ready AI™ 보유하나 경쟁사 대비 투자 열위 (FCF 압박)'],
        ['GenAI 대체 위험', '중위험. XM 통합 플랫폼으로 방어력 상대적 高', '고위험. 분석·보고서 레이어 LLM 즉시 대체 가능 영역 비중 큼'],
        ['FCF 상태', '추정 중 (비공개). 단, 인수 당시 흑자 전환 근접', "'21년~'26년 FCF Coverage 0.40~0.67× / 단 한 번도 1.0× 미달"],
        ['부채 구조', '변동금리 LBO 차입 포함 — 구조 유사 (상세 비공개)', '$1.8B → PIK 누적 $2.78B / SOFR 연동 변동금리'],
        ['Exit 현황', 'IPO 재추진 검토 중 (불확실) / Silver Lake 보유 기간 연장 가능성', 'Exit 실패 확정 — 채권단 경영권 이전'],
        ['Gartner MQ 위치', 'Leader (1위 — Ability to Execute 최상위)', 'Leader (2위 — Completeness of Vision 우수)'],
    ],
    col_widths=[3.0, 7.5, 5.0]
)

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 6 — SaaS M&A 평가 기준(안)
# ════════════════════════════════════════════════════════
add_app_title("【 별첨 6 】 SaaS 기업의 M&A 검토 시 평가 기준(안)")

add_tbl(
    headers=['단계', '체크 항목', 'Medallia 사례 결과', '최소 기준(안)'],
    rows=[
        ['현금창출력', 'FCF Coverage\n(FCF ÷ 예상 이자)', '인수 당시 0.67× — 미달', '≥ 1.5×'],
        ['고객 유지력', 'NRR 지속 가능성 + 업셀링 경로', "IPO 시 126%→116% 하락 추세 / 인수 후 추가 하락 추정 (비공개)", 'NRR ≥ 110%'],
        ['성장+수익 균형', 'Rule of 40', '10 (업계 하위 10%)', '≥ 30'],
        ['GenAI 대체 가능성', '핵심 기능\nGenAI 대체 가능성', '분석 레이어 고위험', '데이터 인프라\n레이어 비중 확인'],
        ['금리 내성', '고금리 3년 지속 시 FCF 시뮬레이션', '시뮬레이션 없이 인수', 'SOFR +2%\n스트레스 테스트'],
        ['EV/Par Coverage', 'EV ÷ Loan Par', "인수 3.56× → '26년 0.79×", '2.0× 이하 경보\n1.5× 이하 위험'],
        ['Exit 계획', 'FCF 기반 구체적 Exit 조건 사전 명시', '"시장 회복 시 IPO" 조건 없음', 'FCF Coverage 2.0× 도달 시 Exit 추진'],
    ],
    col_widths=[2.5, 3.5, 5.5, 4.0]
)

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 7 — PIK 시뮬레이션 & EV/Par Coverage
# ════════════════════════════════════════════════════════
add_app_title('【 별첨 7 】 PIK 원금 누적 시뮬레이션 및 EV/Par Coverage 이탈 과정')

add_app_sub('PIK 원금 누적 시뮬레이션')
add_tbl(
    headers=['기간', 'PIK 비중', '기초 원금', 'PIK이자 누적', '현금이자 지급', '기타 증감', '기말 원금'],
    rows=[
        ["'21~'22", '100%', '$1.80B', '+$0.14B', '$0', '—', '$1.94B'],
        ["'22~'23", '50%',  '$1.94B', '+$0.10B', '$0.10B', '—', '$2.04B'],
        ["'23~'25", '35%',  '$2.04B', '+$0.14B', '$0.18B', '추가차입 +$0.60B', '$2.78B'],
        ["'25~'26", '0%',   '$2.78B', '$0', '$0.27B', '—', '$2.78B → D/E Swap'],
    ],
    col_widths=[1.7, 1.5, 2.0, 2.0, 2.0, 2.5, 3.8]
)
add_fn('PIK(Payment-in-Kind): 현금 대신 원금 증가로 이자를 갚는 구조. 단기 현금 부담을 줄이되 원금이 복리로 증가')

add_app_sub('EV/Par Coverage 임계치 이탈 과정')
add_tbl(
    headers=['시점', '추정 EV', 'Loan Par', 'EV/Par Coverage', '상태 판정'],
    rows=[
        ["'21년 인수", '$6.40B', '$1.80B', '3.56×', '안전'],
        ["'22년 말",   '~$5.20B','$1.94B', '2.68×', '안전'],
        ["'23년 말",   '~$4.00B','$2.20B', '1.82×', '주의 (2.0× 하회)'],
        ["'24년 말",   '~$3.00B','$2.78B', '1.08×', '위험 (1.5× 하회)'],
        ["'26.1분기", '~$2.20B','$2.78B', '0.79×', '임계치 이탈 → 경영권 이전'],
    ],
    col_widths=[2.5, 2.5, 2.5, 3.0, 5.0]
)
add_fn('EV/Par Coverage = EV ÷ Loan Par. 1.0× 미만 = 기업가치가 누적 부채를 충당 못함 → 채권단의 D/E Swap 논리 성립')
add_fn("EV 추정: 당해연도 추정 매출 × 동종 SaaS 평균 EV/Revenue 멀티플 역산 ('21년 9× → '23년 5× → '25년 3×). 상장폐지 이후 추정치 사용.")

add_page_break()


# ════════════════════════════════════════════════════════
# 별첨 8 — GenAI 대체 가능성
# ════════════════════════════════════════════════════════
add_app_title('【 별첨 8 】 SaaS 기업가치 평가의 새로운 기준 : Gen AI 대체 가능성')

add_tbl(
    headers=['항목', '가능성 높음', '가능성 낮음'],
    rows=[
        ['워크플로우\n복잡/자율성', '단일 작업 중심', '부서간 데이터 연계, 권한/예외/연동이 포함된 복합적인 워크플로우'],
        ['데이터의\n독점/폐쇄성', '공개정보 및 인터넷 데이터 기반 범용 기능', '기업 내부에 축적된 프라이빗 데이터를 학습, 산업별 인사이트 제공'],
        ['기능 vs.\n기록 시스템', 'UI/UX의 편리함이나 특정 기능 제공 위주\n(AI 에이전트는 UI의 의미가 희석됨)', '기업의 핵심 정보를 저장/관리, 보안/규제에 의해 통제되는 기록 시스템'],
        ['도메인 전문성/\n맥락 이해도', '범용 LLM/일반 지식으로 해결 가능', '고도의 전문 지식과 업계의 특유 맥락 요구, 환각 발생시 치명적 리스크 존재'],
        ['고객 전환비용/\n생태계 종속성', '타 AI 모델/API로 이전 용이, 대체재 다수', '고객사 내부 시스템과 통합되어 있어 쉽게 들어내기 어려운 구조'],
    ],
    col_widths=[2.8, 5.5, 7.2]
)

add_page_break()


# ════════════════════════════════════════════════════════
# 참고문헌
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
r = p.add_run('참 고 문 헌')
set_font(r, size=14, bold=True)

add_sub_heading('▸ 분석 기사')
add_dash('OnlyCFO, "Thoma Bravo\'s $5B Medallia Loss | The PE Reckoning is Coming..." (\'26.5월) — onlycfo.io')
add_dash('Bloomberg, "Thoma Bravo Declines to Invest More in Troubled Medallia" (\'26.4월) — bloomberg.com')

add_sub_heading('▸ 기업 공식 자료')
add_dash('Medallia 공식 홈페이지 — medallia.com')
add_dash('Thoma Bravo 공식 홈페이지 — thomabravo.com')
add_dash('Thoma Bravo, "Medallia to be Acquired by Thoma Bravo for $6.4 Billion" (보도자료, \'21.7월)')
add_dash('PR Newswire, "Medallia Announces Pricing of Initial Public Offering" (\'19.7월)')

add_sub_heading('▸ 개념 참고')
add_dash('Gartner Magic Quadrant for Voice of the Customer Platforms 2025·2026')
add_dash('Wikipedia, "Medallia" — en.wikipedia.org')
add_dash('Wikipedia, "Thoma Bravo" — en.wikipedia.org')
add_dash('Wikipedia, "Rule of 40" — en.wikipedia.org')
add_dash('McKinsey, "SaaS and the Rule of 40: Keys to the Critical Value Creation Metric"')


# ── 저장 ───────────────────────────────────────────────
output_path = '/home/ubuntu/files/medallia_report_final.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
